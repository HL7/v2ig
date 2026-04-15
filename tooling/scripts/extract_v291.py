#!/usr/bin/env python3
"""
HL7 V2.9.1 Word Document Extraction Script

Parses HL7 V2.9.1 specification Word documents and outputs structured JSON.
Preserves all occurrences (no deduplication) with full provenance metadata.
"""

import json
import os
import re
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from collections import defaultdict

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


# Output directories
OUTPUT_BASE = Path("/workspace/v291-extracted")
SEGMENTS_DIR = OUTPUT_BASE / "segments"
MSG_STRUCTURES_DIR = OUTPUT_BASE / "message-structures"
EVENTS_DIR = OUTPUT_BASE / "events"
DATA_TYPES_DIR = OUTPUT_BASE / "data-types" / "complex"
QUERY_PARAMS_DIR = OUTPUT_BASE / "query-parameters"

# Source directory
SOURCE_DIR = Path("/workspace/v2plus_docx")

# Files to skip
SKIP_FILES = {"CH02C_Tables.docx", "CH01_Intro.docx"}

# Global warning/error log
extraction_log = []


def log_warning(message: str):
    """Log a warning message."""
    print(f"⚠️  {message}")
    extraction_log.append({"type": "warning", "message": message})


def log_error(message: str):
    """Log an error message."""
    print(f"❌ {message}")
    extraction_log.append({"type": "error", "message": message})


def log_info(message: str):
    """Log an info message."""
    print(f"ℹ️  {message}")


def normalize_header(header: str) -> str:
    """Normalize table header text for comparison."""
    # Remove extra whitespace, convert to uppercase
    return re.sub(r'\s+', ' ', header.strip().upper())


def fuzzy_match_headers(row_cells: List[str], expected_patterns: List[str]) -> bool:
    """
    Check if table headers fuzzy match expected patterns.

    Args:
        row_cells: List of cell text from the header row
        expected_patterns: List of substrings that should appear in headers

    Returns:
        True if enough patterns match
    """
    normalized = [normalize_header(cell) for cell in row_cells]
    combined = ' '.join(normalized)

    matches = sum(1 for pattern in expected_patterns if pattern.upper() in combined)
    return matches >= len(expected_patterns) * 0.7  # 70% match threshold


def extract_chapter_number(filename: str) -> str:
    """Extract chapter number from filename."""
    match = re.match(r'CH(\d+[A-Z]?)', filename)
    return match.group(1) if match else "unknown"


def get_table_cells_text(table: Table) -> List[List[str]]:
    """Extract text from all cells in a table."""
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def identify_table_type(table: Table, preceding_style: Optional[str] = None) -> Optional[str]:
    """
    Identify the type of table based on headers and context.

    Returns:
        One of: 'segment_fields', 'message_structure', 'ack_choreography',
                'data_type_components', 'query_parameters', or None
    """
    if len(table.rows) == 0:
        return None

    first_row = [cell.text.strip() for cell in table.rows[0].cells]
    num_cols = len(first_row)

    # Check for data type components first (9 columns with specific style)
    # This must come before segment fields because both are 9 columns
    if num_cols == 9 and preceding_style == 'Component Table Caption':
        component_patterns = ['SEQ', 'COMPONENT', 'COMMENTS']
        if fuzzy_match_headers(first_row, component_patterns):
            return 'data_type_components'

    # Check for segment field table
    # Standard: 9 columns. But some tables have duplicate columns (18, 13, etc.)
    # due to Word formatting artifacts. Use header content to identify.
    segment_patterns = ['SEQ', 'ELEMENT', 'ITEM']
    if fuzzy_match_headers(first_row, segment_patterns):
        return 'segment_fields'

    # Check for message structure table (typically 4-5 cols, but duplicates
    # can produce 6+). Use header matching instead of strict column count.
    if num_cols >= 4:
        msg_patterns = ['SEGMENTS', 'DESCRIPTION', 'STATUS', 'CHAPTER']
        if fuzzy_match_headers(first_row, msg_patterns):
            return 'message_structure'

    # Check for acknowledgment choreography (4-6 columns).
    # Handle both "Acknowledgment" and "Acknowledgement" spellings.
    if num_cols in [4, 5, 6]:
        combined = ' '.join(first_row).upper()
        if 'ACKNOWLEDG' in combined and 'CHOREOGRAPHY' in combined:
            return 'ack_choreography'

    # Check for query parameters (13 columns)
    if num_cols == 13:
        query_patterns = ['FIELD', 'KEY', 'SEARCH', 'SORT', 'MATCH']
        if fuzzy_match_headers(first_row, query_patterns):
            return 'query_parameters'

    return None


def deduplicate_segment_columns(headers: List[str], row: List[str]) -> List[str]:
    """
    Deduplicate columns in a segment field table row.

    Some Word tables have duplicate columns (18 or 13 cols instead of 9) due
    to formatting artifacts. This function finds the first occurrence of each
    canonical column and returns a normalized 9-element row.

    Canonical order: SEQ, LEN, C.LEN, DT, OPT, RP/#, TBL#, ITEM#, ELEMENT NAME
    """
    canonical = ['SEQ', 'LEN', 'C.LEN', 'DT', 'OPT', 'RP/#', 'TBL#', 'ITEM', 'ELEMENT']
    normalized_headers = [normalize_header(h) for h in headers]

    indices = []
    used = set()
    for pattern in canonical:
        for i, h in enumerate(normalized_headers):
            if i not in used and pattern in h:
                indices.append(i)
                used.add(i)
                break
        else:
            # Pattern not found — append None placeholder
            indices.append(None)

    result = []
    for idx in indices:
        if idx is not None and idx < len(row):
            result.append(row[idx])
        else:
            result.append("")
    return result


def parse_segment_field_table(table: Table) -> List[Dict[str, Any]]:
    """
    Parse a segment field definition table.

    Standard format: 9 columns (SEQ | LEN | C.LEN | DT | OPT | RP/# | TBL# | ITEM# | ELEMENT NAME).
    Also handles tables with duplicate columns (18, 13 cols) by deduplicating.
    """
    cells = get_table_cells_text(table)
    if len(cells) < 2:
        return []

    headers = cells[0]
    num_cols = len(headers)
    needs_dedup = num_cols != 9

    fields = []
    # Skip header row
    for row in cells[1:]:
        if len(row) < min(9, num_cols):
            continue

        # Deduplicate if needed
        if needs_dedup:
            row = deduplicate_segment_columns(headers, row)

        # Skip empty rows or separator rows
        if not row[0].strip() or row[0].strip() == '-':
            continue

        field = {
            "sequence": row[0].strip(),
            "length": row[1].strip(),
            "confLength": row[2].strip(),
            "dataType": row[3].strip(),
            "optionality": row[4].strip(),
            "repetition": row[5].strip(),
            "tableBinding": row[6].strip(),
            "itemNumber": row[7].strip(),
            "name": row[8].strip()
        }
        fields.append(field)

    return fields


def parse_message_structure_table(table: Table) -> Tuple[List[Dict], List[Dict]]:
    """
    Parse a message structure table.

    Returns:
        Tuple of (raw_rows, parsed_structure)
    """
    cells = get_table_cells_text(table)
    if len(cells) < 2:
        return [], []

    raw_rows = []
    parsed_structure = []

    # Skip header row
    for row_idx, row in enumerate(cells[1:], start=1):
        if len(row) < 4:
            continue

        # Handle both 4 and 5 column tables (sometimes Chapter appears twice)
        segments_col = row[0].strip()
        description_col = row[1].strip()
        status_col = row[2].strip() if len(row) > 2 else ""
        chapter_col = row[3].strip() if len(row) > 3 else ""

        # Skip empty rows
        if not segments_col:
            continue

        raw_row = {
            "segments": segments_col,
            "description": description_col,
            "status": status_col,
            "chapter": chapter_col
        }
        raw_rows.append(raw_row)

        # Parse the segments column for structure
        try:
            parsed = parse_segment_notation(segments_col, description_col, chapter_col)
            parsed_structure.append(parsed)
        except Exception as e:
            log_warning(f"Failed to parse segment notation '{segments_col}': {e}")
            parsed_structure.append({
                "type": "unparsed",
                "raw": segments_col,
                "description": description_col,
                "parseError": str(e)
            })

    return raw_rows, parsed_structure


def _parse_table_no_header_skip(table: Table) -> Tuple[List[Dict], List[Dict]]:
    """Parse a message structure continuation table without skipping the first row."""
    cells = get_table_cells_text(table)
    raw_rows = []
    parsed_structure = []

    for row in cells:  # No skip — all rows are data
        if len(row) < 4:
            continue

        segments_col = row[0].strip()
        description_col = row[1].strip()
        status_col = row[2].strip() if len(row) > 2 else ""
        chapter_col = row[3].strip() if len(row) > 3 else ""

        if not segments_col:
            continue

        raw_row = {
            "segments": segments_col,
            "description": description_col,
            "status": status_col,
            "chapter": chapter_col
        }
        raw_rows.append(raw_row)

        try:
            parsed = parse_segment_notation(segments_col, description_col, chapter_col)
            parsed_structure.append(parsed)
        except Exception as e:
            log_warning(f"Failed to parse segment notation '{segments_col}': {e}")
            parsed_structure.append({
                "type": "unparsed",
                "raw": segments_col,
                "description": description_col,
                "parseError": str(e)
            })

    return raw_rows, parsed_structure


def parse_segment_notation(notation: str, description: str, chapter: str) -> Dict[str, Any]:
    """
    Parse segment notation to determine optionality and repetition.

    Notation:
    - SEG = required, no repeat
    - [ SEG ] = optional
    - { SEG } = required, repeating
    - [{ SEG }] or [ { SEG } ] = optional, repeating
    - --- PREFIX = group boundary marker
    - < = choice group begin
    - | = choice alternative separator
    - > = choice group end
    """
    original = notation
    notation = notation.strip()

    # Check for choice group markers
    if notation == '<':
        return {
            "type": "choice_start",
            "description": description
        }
    if notation == '>':
        return {
            "type": "choice_end",
            "description": description
        }
    if notation == '|':
        return {
            "type": "choice_alternative",
            "description": description
        }

    # Check for group boundary markers
    if notation.startswith('---'):
        # Extract group name
        group_name = notation[3:].strip()
        if group_name.endswith('begin'):
            return {
                "type": "group_start",
                "name": group_name.replace('begin', '').strip(),
                "description": description
            }
        elif group_name.endswith('end'):
            return {
                "type": "group_end",
                "name": group_name.replace('end', '').strip(),
                "description": description
            }
        else:
            return {
                "type": "group_marker",
                "name": group_name,
                "description": description
            }

    # Parse optionality and repetition from brackets
    optional = False
    repeating = False

    # Check for outer brackets (optionality)
    if notation.startswith('[') and notation.endswith(']'):
        optional = True
        notation = notation[1:-1].strip()

    # Check for inner braces (repetition)
    if notation.startswith('{') and notation.endswith('}'):
        repeating = True
        notation = notation[1:-1].strip()

    # If there are still brackets, might be [{ SEG }] notation
    if '[' in notation or '{' in notation:
        # Try to extract the segment code
        segment_code = re.sub(r'[\[\]{}]', '', notation).strip()
        # Re-check the original for optionality/repetition
        if '[' in original:
            optional = True
        if '{' in original:
            repeating = True
    else:
        segment_code = notation

    # Determine optionality flag
    opt_flag = 'O' if optional else 'R'

    return {
        "type": "segment",
        "code": segment_code,
        "description": description,
        "optionality": opt_flag,
        "repetition": repeating,
        "chapter": chapter
    }


def expand_event_codes(raw: str) -> List[str]:
    """
    Expand event code expressions into individual event codes.

    Examples:
        "A01" → ["A01"]
        "S01-S11" → ["S01", "S02", ..., "S11"]
        "S12-S24,S26,S27" → ["S12", "S13", ..., "S24", "S26", "S27"]
        "PC1-PCH,PCJ,PCK,PCL" → ["PC1", ..., "PCH", "PCJ", "PCK", "PCL"]
    """
    raw = raw.strip()

    # Split on commas first
    parts = [p.strip() for p in raw.split(',')]
    result = []

    for part in parts:
        if '-' in part:
            # It's a range like "S01-S11" or "PC1-PCH"
            range_parts = part.split('-', 1)
            if len(range_parts) == 2:
                start, end = range_parts[0].strip(), range_parts[1].strip()
                expanded = expand_event_range(start, end)
                result.extend(expanded)
            else:
                result.append(part)
        else:
            result.append(part)

    return result


def expand_event_range(start: str, end: str) -> List[str]:
    """
    Expand a range like S01-S11 or PC1-PCH into individual codes.

    Handles numeric suffixes (S01-S11) and alpha suffixes (PC1-PCH).
    """
    # Find common prefix
    prefix_len = 0
    for i in range(min(len(start), len(end))):
        if start[i] == end[i]:
            prefix_len = i + 1
        else:
            break

    # Handle case where prefix is the letter part (e.g., "S" for S01-S11)
    prefix = start[:prefix_len]
    start_suffix = start[prefix_len:]
    end_suffix = end[prefix_len:]

    # If no common prefix found, try single-char prefix
    if not prefix:
        return [start, end]

    # Try numeric range
    if start_suffix.isdigit() and end_suffix.isdigit():
        width = len(start_suffix)
        codes = []
        for i in range(int(start_suffix), int(end_suffix) + 1):
            codes.append(f"{prefix}{str(i).zfill(width)}")
        return codes

    # Try alphanumeric range (e.g., PC1...PCH or PCA...PCL)
    # Use character ordinals for single-char suffixes, skipping non-alphanumeric
    if len(start_suffix) == 1 and len(end_suffix) == 1:
        codes = []
        for c in range(ord(start_suffix), ord(end_suffix) + 1):
            ch = chr(c)
            if ch.isalnum():
                codes.append(f"{prefix}{ch}")
        return codes

    # Fallback: just return start and end
    return [start, end]


def parse_ack_choreography_table(table: Table) -> Dict[str, Any]:
    """
    Parse an acknowledgment choreography table.

    Expected structure:
    - Row 0: Merged header with "Acknowledgment Choreography"
    - Row 1: Message identifier repeated
    - Row 2: Column labels
    - Row 3: MSH-15/MSH.15 values
    - Row 4: MSH-16/MSH.16 values
    - Row 5: Immediate Ack values
    - Row 6: Application Ack values
    """
    cells = get_table_cells_text(table)

    if len(cells) < 6:
        log_warning(f"Ack choreography table has only {len(cells)} rows, expected at least 6")
        return {}

    # Extract message identifier from row 1
    msg_identifier = cells[1][0].strip() if len(cells[1]) > 0 else ""

    # Extract column labels from row 2
    columns = [cell.strip() for cell in cells[2]]

    # Extract field values from subsequent rows
    msh15 = [cell.strip() for cell in cells[3]] if len(cells) > 3 else []
    msh16 = [cell.strip() for cell in cells[4]] if len(cells) > 4 else []
    immediate_ack = [cell.strip() for cell in cells[5]] if len(cells) > 5 else []
    application_ack = [cell.strip() for cell in cells[6]] if len(cells) > 6 else []

    return {
        "messageIdentifier": msg_identifier,
        "columns": columns,
        "msh15": msh15,
        "msh16": msh16,
        "immediateAck": immediate_ack,
        "applicationAck": application_ack
    }


def parse_data_type_components_table(table: Table) -> List[Dict[str, Any]]:
    """
    Parse a data type components table (9 columns).

    Expected columns: SEQ | LEN | C.LEN | DT | OPT | TBL# | COMPONENT NAME | COMMENTS | SEC.REF.
    """
    cells = get_table_cells_text(table)
    if len(cells) < 2:
        return []

    components = []
    # Skip header row
    for row in cells[1:]:
        if len(row) < 9:
            continue

        # Skip empty rows
        if not row[0].strip():
            continue

        component = {
            "sequence": row[0].strip(),
            "length": row[1].strip(),
            "confLength": row[2].strip(),
            "dataType": row[3].strip(),
            "optionality": row[4].strip(),
            "tableBinding": row[5].strip(),
            "name": row[6].strip(),
            "comments": row[7].strip(),
            "sectionRef": row[8].strip()
        }
        components.append(component)

    return components


def parse_query_parameters_table(table: Table) -> List[Dict[str, Any]]:
    """Parse a query parameters table (13 columns)."""
    cells = get_table_cells_text(table)
    if len(cells) < 2:
        return []

    params = []
    for row in cells[1:]:
        if len(row) < 13:
            continue

        if not row[0].strip():
            continue

        param = {
            "fieldSeq": row[0].strip(),
            "fieldName": row[1].strip(),
            "keySearch": row[2].strip(),
            "sort": row[3].strip(),
            "length": row[4].strip(),
            "type": row[5].strip(),
            "optionality": row[6].strip(),
            "repetition": row[7].strip(),
            "matchOp": row[8].strip(),
            "table": row[9].strip(),
            "segmentFieldName": row[10].strip(),
            "loincOrHl7Code": row[11].strip(),
            "elementName": row[12].strip()
        }
        params.append(param)

    return params


def extract_segment_code_from_heading(heading: str) -> Optional[str]:
    """
    Extract segment code from heading text.

    Examples:
    - "MSH – Message Header Segment" -> "MSH"
    - "PID - Patient Identification Segment" -> "PID"
    - "HL7 Attribute Table - EVN – Event Type" -> "EVN"
    - "HL7 Attribute Table - MSH - Message Header" -> "MSH"
    """
    # First try the "HL7 Attribute Table - CODE –" pattern
    match = re.search(r'HL7\s+Attribute\s+Table\s*[-–—]\s*([A-Z0-9]{2,4})\s*[-–—]', heading, re.IGNORECASE)
    if match:
        return match.group(1)

    # Otherwise look for 2-4 letter code before dash at start of line
    match = re.match(r'^([A-Z0-9]{2,4})\s*[-–—]\s*', heading)
    return match.group(1) if match else None


def extract_segment_name_from_heading(heading: str) -> str:
    """
    Extract the human-readable segment name from a heading.

    Examples:
    - "MSH – Message Header Segment" -> "Message Header Segment"
    - "HL7 Attribute Table - PV1 - Patient Visit" -> "Patient Visit"
    - "HL7 Attribute Table - IAM - Patient Adverse Reaction Information" -> "Patient Adverse Reaction Information"
    - "PID - Patient Identification Segment" -> "Patient Identification Segment"
    """
    # For "HL7 Attribute Table - CODE - Name" pattern
    match = re.search(r'HL7\s+Attribute\s+Table\s*[-–—]\s*[A-Z0-9]{2,4}\s*[-–—]\s*(.+)', heading, re.IGNORECASE)
    if match:
        return match.group(1).strip()

    # For "CODE – Name" or "CODE - Name" pattern
    # Split on any dash type, take everything after the first dash following the code
    match = re.match(r'^[A-Z0-9]{2,4}\s*[-–—]\s*(.+)', heading)
    if match:
        return match.group(1).strip()

    return ""


def extract_data_type_code_from_caption(caption: str) -> Optional[str]:
    """
    Extract data type code from caption.

    Examples:
    - "HL7 Component Table - AD – Address" -> "AD"
    - "HL7 Component Table – CCP - Channel Calibration Parameters" -> "CCP"
    - "HL7 Component Table – XPN – Extended Person Name" -> "XPN"
    """
    # Look for pattern like "- CODE –", "– CODE -", or "– CODE –"
    # Handle all combinations of single dash, en-dash, em-dash
    match = re.search(r'[-–—]\s*([A-Z]{2,10})\s*[-–—]', caption)
    return match.group(1) if match else None


def extract_data_type_name_from_caption(caption: str) -> str:
    """
    Extract the human-readable data type name from a caption.

    Examples:
    - "HL7 Component Table - AD – Address" -> "Address"
    - "HL7 Component Table – XPN – Extended Person Name" -> "Extended Person Name"
    """
    # Find the data type code, then take everything after the next dash
    match = re.search(r'[-–—]\s*[A-Z]{2,10}\s*[-–—]\s*(.+)', caption)
    if match:
        return match.group(1).strip()
    return ""


def find_preceding_heading(doc: Document, table_obj: Table, para_cache: Dict = None) -> Tuple[Optional[str], Optional[str]]:
    """
    Find the heading that precedes a table.

    Args:
        doc: The Word document
        table_obj: The table to find the heading for
        para_cache: Optional cache of element -> paragraph mappings

    Returns:
        Tuple of (heading_text, style_name)
    """
    # Build paragraph cache if not provided
    if para_cache is None:
        para_cache = {}
        for p in doc.paragraphs:
            para_cache[id(p._element)] = p

    # Get the target element directly
    target_element = table_obj._element

    # Find the position of this table in the body
    recent_heading = None
    recent_heading_style = None

    for element in doc.element.body:
        if element.tag.endswith('p'):
            # Look up the paragraph from cache
            para = para_cache.get(id(element))

            if para:
                style = para.style.name if para.style else None
                text = para.text.strip()

                # Track headings and captions
                if style and ('Heading' in style or 'Caption' in style) and text:
                    recent_heading = text
                    recent_heading_style = style

        elif element.tag.endswith('tbl'):
            # This is a table - check if it's our target
            if element is target_element:
                # Found our table - return the most recent heading
                return recent_heading, recent_heading_style

    return None, None


def build_element_map(doc: Document) -> Dict:
    """
    Build a mapping of document structure for fast lookups.

    Returns a dict with:
      - 'para_map': element_id -> paragraph object
      - 'table_map': element_id -> table object
      - 'headings_before_table': table_element_id -> (heading_text, style_name)
    """
    # Build paragraph and table maps
    para_map = {}
    for p in doc.paragraphs:
        para_map[id(p._element)] = p

    table_map = {}
    for t in doc.tables:
        table_map[id(t._element)] = t

    # Build heading-before-table map
    headings_before_table = {}
    recent_heading = None
    recent_heading_style = None

    for element in doc.element.body:
        if element.tag.endswith('p'):
            para = para_map.get(id(element))
            if para:
                style = para.style.name if para.style else None
                text = para.text.strip()

                if style and ('Heading' in style or 'Caption' in style) and text:
                    recent_heading = text
                    recent_heading_style = style

        elif element.tag.endswith('tbl'):
            # Record the most recent heading for this table
            headings_before_table[id(element)] = (recent_heading, recent_heading_style)

    return {
        'para_map': para_map,
        'table_map': table_map,
        'headings_before_table': headings_before_table
    }


def process_document(filepath: Path) -> Dict[str, Any]:
    """
    Process a single Word document and extract all structured data.

    Returns:
        Dictionary with extracted data organized by type
    """
    filename = filepath.name
    chapter = extract_chapter_number(filename)

    log_info(f"Processing {filename} (Chapter {chapter})")

    doc = Document(filepath)

    # Build element map once for the entire document
    element_map = build_element_map(doc)

    # Storage for extracted data
    segments = defaultdict(list)  # segment_code -> list of occurrences
    message_structures = []
    events = []
    data_types = defaultdict(list)  # type_code -> list of occurrences
    query_params = []

    # Track the last caption for context
    last_caption = None
    last_caption_style = None

    # Track tables consumed as continuations (skip in main loop)
    consumed_tables = set()

    # Iterate through tables
    for table_idx, table in enumerate(doc.tables):
        if table_idx in consumed_tables:
            continue
        # Find preceding heading/caption from pre-built map
        heading_text, heading_style = element_map['headings_before_table'].get(id(table._element), (None, None))

        # Update last caption if we found one
        if heading_style and 'Caption' in heading_style:
            last_caption = heading_text
            last_caption_style = heading_style

        # Identify table type
        table_type = identify_table_type(table, last_caption_style)

        if table_type is None:
            continue

        # Create provenance metadata
        provenance = {
            "sourceFile": filename,
            "chapter": chapter,
            "sectionHeading": heading_text or "",
            "tableIndex": table_idx,
            "captionText": last_caption or ""
        }

        # Process based on type
        if table_type == 'segment_fields':
            # Extract segment code from heading
            segment_code = extract_segment_code_from_heading(heading_text) if heading_text else None

            if not segment_code:
                log_warning(f"Could not extract segment code from heading '{heading_text}' at table {table_idx}")
                continue

            fields = parse_segment_field_table(table)
            if fields:
                segments[segment_code].append({
                    "name": extract_segment_name_from_heading(heading_text),
                    "provenance": provenance,
                    "fields": fields
                })
                log_info(f"  → Extracted segment {segment_code} with {len(fields)} fields")

        elif table_type == 'message_structure':
            raw_rows, parsed_structure = parse_message_structure_table(table)

            # Extract structure ID from caption
            structure_id = None
            if last_caption:
                # Look for pattern like "ADT^A01^ADT_A01" or "ACK^varies^ACK"
                match = re.search(r'([A-Za-z0-9]+)\^([A-Za-z0-9,\-]+)\^([A-Za-z0-9_]+)', last_caption)
                if match:
                    raw_structure = match.group(3)
                    # Skip "varies" — use message type as structure ID instead
                    if raw_structure.lower() == 'varies':
                        structure_id = match.group(1).upper()
                    else:
                        structure_id = raw_structure

            # Check for continuation tables (split tables from Word formatting)
            # Continuation tables have the same column format but no header row
            cont_idx = table_idx + 1
            while cont_idx < len(doc.tables):
                cont_table = doc.tables[cont_idx]
                if len(cont_table.rows) == 0:
                    cont_idx += 1
                    continue

                cont_first = [cell.text.strip() for cell in cont_table.rows[0].cells]
                cont_cols = len(cont_first)

                # Stop if this is an ack choreography table
                cont_combined = ' '.join(cont_first).upper()
                if 'ACKNOWLEDG' in cont_combined and 'CHOREOGRAPHY' in cont_combined:
                    break

                # Stop if this has a header row (new message structure table)
                if cont_cols >= 4 and cont_first[0].upper().startswith('SEGMENT'):
                    break

                # Stop if column count doesn't match message structure format
                if cont_cols < 4:
                    break

                # Stop if it looks like a different table type
                cont_type = identify_table_type(cont_table, None)
                if cont_type and cont_type != 'message_structure':
                    break

                # This is a continuation — merge its rows
                cont_rows, cont_parsed = parse_message_structure_table(cont_table)
                if cont_rows:
                    # Continuation tables have no header, but parse_message_structure_table
                    # skips the first row as header. Check if it was a real header.
                    # If the first cell doesn't look like a header, re-parse without skip.
                    if cont_first[0].upper() not in ('SEGMENTS', 'SEGMENT', ''):
                        # Re-parse treating ALL rows as data (no header skip)
                        cont_rows_full, cont_parsed_full = _parse_table_no_header_skip(cont_table)
                        raw_rows.extend(cont_rows_full)
                        parsed_structure.extend(cont_parsed_full)
                    else:
                        raw_rows.extend(cont_rows)
                        parsed_structure.extend(cont_parsed)
                    consumed_tables.add(cont_idx)
                    log_info(f"  → Merged continuation table {cont_idx} ({len(cont_rows)} rows)")
                else:
                    # Empty continuation — still consume it
                    consumed_tables.add(cont_idx)

                cont_idx += 1

            if raw_rows:
                message_structures.append({
                    "structureId": structure_id or "UNKNOWN",
                    "caption": last_caption or "",
                    "provenance": provenance,
                    "rawRows": raw_rows,
                    "parsedStructure": parsed_structure
                })
                log_info(f"  → Extracted message structure {structure_id} with {len(raw_rows)} rows")

        elif table_type == 'ack_choreography':
            ack_data = parse_ack_choreography_table(table)
            if ack_data:
                # Extract event code from message identifier
                # Handles: ADT^A01^ADT_A01, ACK^varies^ACK, ranges like S01-S11
                msg_id = ack_data.get('messageIdentifier', '')
                match = re.match(r'([A-Za-z]+)\^([A-Za-z0-9,\-]+)\^', msg_id)
                if match:
                    message_type = match.group(1).upper()
                    raw_event = match.group(2)

                    # Skip "varies" — not a real event code
                    if raw_event.lower() == 'varies':
                        log_info(f"  → Skipping 'varies' event in {msg_id}")
                    else:
                        # Expand event code ranges (e.g., "S01-S11" → S01..S11)
                        expanded_events = expand_event_codes(raw_event)
                        parts = msg_id.split('^')
                        structure_id = parts[2] if len(parts) > 2 else ""

                        for event_code in expanded_events:
                            events.append({
                                "eventCode": event_code,
                                "messageType": message_type,
                                "name": heading_text or "",
                                "structureId": structure_id,
                                "provenance": provenance,
                                "acknowledgmentChoreography": ack_data
                            })
                        log_info(f"  → Extracted {len(expanded_events)} event(s) from {msg_id}")

        elif table_type == 'data_type_components':
            # Extract data type code from caption
            dt_code = extract_data_type_code_from_caption(last_caption) if last_caption else None

            if not dt_code:
                log_warning(f"Could not extract data type code from caption '{last_caption}' at table {table_idx}")
                continue

            components = parse_data_type_components_table(table)
            if components:
                data_types[dt_code].append({
                    "name": extract_data_type_name_from_caption(last_caption),
                    "caption": last_caption or "",
                    "provenance": provenance,
                    "components": components
                })
                log_info(f"  → Extracted data type {dt_code} with {len(components)} components")

        elif table_type == 'query_parameters':
            params = parse_query_parameters_table(table)
            if params:
                query_params.append({
                    "provenance": provenance,
                    "parameters": params
                })
                log_info(f"  → Extracted query parameters with {len(params)} params")

    # Derive events from message structure captions for chapters without
    # ack choreography tables. This fills the gap for CH04, CH05, CH07, etc.
    # Also handles multi-event captions like PGL^PC6-PC8^PGL_PC6.
    ack_event_codes = {e["eventCode"] for e in events}
    for ms in message_structures:
        caption = ms.get("caption", "")
        match = re.search(r'([A-Za-z0-9]+)\^([A-Za-z0-9,\-]+)\^([A-Za-z0-9_]+)', caption)
        if match:
            raw_event = match.group(2)
            message_type = match.group(1).upper()
            structure_id = match.group(3)

            if raw_event.lower() == 'varies':
                continue

            # Expand event code ranges (e.g., PC6-PC8 → PC6, PC7, PC8)
            expanded = expand_event_codes(raw_event.upper())

            for event_code in expanded:
                if event_code in ack_event_codes:
                    continue

                ack_event_codes.add(event_code)
                events.append({
                    "eventCode": event_code,
                    "messageType": message_type,
                    "name": caption,
                    "structureId": structure_id,
                    "provenance": ms["provenance"],
                    "acknowledgmentChoreography": None,
                    "derivedFrom": "message_structure_caption"
                })

            if expanded:
                log_info(f"  → Derived {len(expanded)} event(s) from structure caption: {caption}")

    return {
        "segments": segments,
        "message_structures": message_structures,
        "events": events,
        "data_types": data_types,
        "query_params": query_params
    }


def write_segment_files(all_segments: Dict[str, List[Dict]]):
    """Write segment data to JSON files."""
    SEGMENTS_DIR.mkdir(parents=True, exist_ok=True)

    for code, occurrences in all_segments.items():
        output_file = SEGMENTS_DIR / f"{code}.json"
        data = {
            "code": code,
            "occurrences": occurrences
        }
        with open(output_file, 'w') as f:
            json.dump(data, f, indent=2)

        log_info(f"Wrote {output_file} ({len(occurrences)} occurrence(s))")


def write_message_structure_files(all_structures: List[Dict]):
    """Write message structure data to JSON files."""
    MSG_STRUCTURES_DIR.mkdir(parents=True, exist_ok=True)

    for idx, structure in enumerate(all_structures):
        structure_id = structure['structureId']
        chapter = structure['provenance']['chapter']

        # Create unique filename with index to preserve ALL occurrences
        filename = f"{structure_id}_{chapter}_{idx}.json"
        output_file = MSG_STRUCTURES_DIR / filename

        with open(output_file, 'w') as f:
            json.dump(structure, f, indent=2)

        log_info(f"Wrote {output_file}")


def write_event_files(all_events: List[Dict]):
    """Write event data to JSON files."""
    EVENTS_DIR.mkdir(parents=True, exist_ok=True)

    for idx, event in enumerate(all_events):
        event_code = event['eventCode']
        chapter = event['provenance']['chapter']

        # Create unique filename with index to preserve ALL occurrences
        filename = f"{event_code}_{chapter}_{idx}.json"
        output_file = EVENTS_DIR / filename

        with open(output_file, 'w') as f:
            json.dump(event, f, indent=2)

        log_info(f"Wrote {output_file}")


def write_data_type_files(all_data_types: Dict[str, List[Dict]]):
    """Write data type data to JSON files."""
    DATA_TYPES_DIR.mkdir(parents=True, exist_ok=True)

    for code, occurrences in all_data_types.items():
        output_file = DATA_TYPES_DIR / f"{code}.json"
        data = {
            "code": code,
            "occurrences": occurrences
        }
        with open(output_file, 'w') as f:
            json.dump(data, f, indent=2)

        log_info(f"Wrote {output_file} ({len(occurrences)} occurrence(s))")


def write_query_param_files(all_query_params: List[Dict]):
    """Write query parameter data to JSON files."""
    if not all_query_params:
        return

    QUERY_PARAMS_DIR.mkdir(parents=True, exist_ok=True)

    for idx, query_data in enumerate(all_query_params):
        chapter = query_data['provenance']['chapter']
        filename = f"query_params_{chapter}_{idx}.json"
        output_file = QUERY_PARAMS_DIR / filename

        with open(output_file, 'w') as f:
            json.dump(query_data, f, indent=2)

        log_info(f"Wrote {output_file}")


def main():
    """Main extraction process."""
    print("=" * 80)
    print("HL7 V2.9.1 Word Document Extraction")
    print("=" * 80)
    print()

    # Get all .docx files to process
    all_files = sorted(SOURCE_DIR.glob("*.docx"))
    files_to_process = [f for f in all_files if f.name not in SKIP_FILES]

    print(f"Found {len(all_files)} .docx files, processing {len(files_to_process)}")
    print()

    # Aggregate storage
    all_segments = defaultdict(list)
    all_message_structures = []
    all_events = []
    all_data_types = defaultdict(list)
    all_query_params = []

    # Process each document
    for filepath in files_to_process:
        try:
            result = process_document(filepath)

            # Merge results
            for code, occurrences in result['segments'].items():
                all_segments[code].extend(occurrences)

            all_message_structures.extend(result['message_structures'])
            all_events.extend(result['events'])

            for code, occurrences in result['data_types'].items():
                all_data_types[code].extend(occurrences)

            all_query_params.extend(result['query_params'])

            print()

        except Exception as e:
            log_error(f"Failed to process {filepath.name}: {e}")
            import traceback
            traceback.print_exc()
            print()

    # Write output files
    print("=" * 80)
    print("Writing output files...")
    print("=" * 80)
    print()

    write_segment_files(all_segments)
    write_message_structure_files(all_message_structures)
    write_event_files(all_events)
    write_data_type_files(all_data_types)
    write_query_param_files(all_query_params)

    # Write metadata summary
    metadata = {
        "extractionDate": "2026-02-26",
        "sourceDirectory": str(SOURCE_DIR),
        "filesProcessed": [f.name for f in files_to_process],
        "counts": {
            "segments": len(all_segments),
            "segmentOccurrences": sum(len(occs) for occs in all_segments.values()),
            "messageStructures": len(all_message_structures),
            "events": len(all_events),
            "dataTypes": len(all_data_types),
            "dataTypeOccurrences": sum(len(occs) for occs in all_data_types.values()),
            "queryParameterSets": len(all_query_params)
        },
        "log": extraction_log
    }

    metadata_file = OUTPUT_BASE / "metadata.json"
    with open(metadata_file, 'w') as f:
        json.dump(metadata, f, indent=2)

    print()
    print("=" * 80)
    print("Extraction Summary")
    print("=" * 80)
    print(f"Segments: {metadata['counts']['segments']} unique codes, {metadata['counts']['segmentOccurrences']} total occurrences")
    print(f"Message Structures: {metadata['counts']['messageStructures']}")
    print(f"Events: {metadata['counts']['events']}")
    print(f"Data Types: {metadata['counts']['dataTypes']} unique codes, {metadata['counts']['dataTypeOccurrences']} total occurrences")
    print(f"Query Parameter Sets: {metadata['counts']['queryParameterSets']}")
    print()
    print(f"Warnings: {len([log for log in extraction_log if log['type'] == 'warning'])}")
    print(f"Errors: {len([log for log in extraction_log if log['type'] == 'error'])}")
    print()
    print(f"Metadata written to: {metadata_file}")
    print()
    print("✅ Extraction complete!")


if __name__ == '__main__':
    main()
