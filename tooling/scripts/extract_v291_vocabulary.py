#!/usr/bin/env python3
"""
Extract vocabulary/code tables from V2.9.1 CH02C_Tables.docx.

Extracts all 797 code tables, each containing up to 7 sub-sections:
  1. Concept Domain Information
  2. Code System Identification Information
  3. Code System Version Information
  4. Value Set Information
  5. Binding Information
  6. Table Metadata
  7. Coded Content (Value, Display Name, Definition, Comment, Status)

Output: One JSON file per table in v291-extracted/vocabulary/
        Plus a summary index at v291-extracted/vocabulary-index.json
        Plus a deviation log at v291-extracted/vocabulary-deviations.json

FIDELITY POLICY
---------------
The published Chapter 2C is the source. This extractor is deliberately
conservative about "cleaning" it, because any silent cleanup is a divergence
the reviewer never gets to see.

TWO normalizations are applied automatically, both defined in
`vocabulary_text_policy.py` and both recorded in the deviation log:

  1. Leading/trailing whitespace is stripped from every cell.   (ADR-0008 D2)
  2. In descriptive prose only, a run of two or more spaces following a
     period is collapsed to a single space.                     (ADR-0008 D3)

Everything else is PRESERVED exactly as published, and merely *reported*:
  * internal double spaces that do NOT follow a period
    (e.g. "Code system of concepts  which specify...")
  * non-breaking spaces
  * newlines inside a cell (usually legitimate paragraph structure)
  * en/em dashes and typographic quotes

Anything the extractor cannot parse, or chooses to skip, is also recorded
rather than dropped silently -- see `sourceIssues` on each table record.

Usage:
    python3 tooling/scripts/extract_v291_vocabulary.py
"""

import datetime
import json
import os
import sys
import re
from pathlib import Path

try:
    import docx
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

sys.path.insert(0, str(Path(__file__).resolve().parent))
from vocabulary_text_policy import (  # noqa: E402
    DESCRIPTIVE_FIELDS,
    normalize_descriptive_text,
)


DOCX_PATH = "v2plus_docx/CH02C_Tables.docx"
OUTPUT_DIR = "v291-extracted/vocabulary"
INDEX_PATH = "v291-extracted/vocabulary-index.json"
DEVIATIONS_PATH = "v291-extracted/vocabulary-deviations.json"

NBSP = " "


class DeviationLog:
    """Collects every place the extracted value differs from the published text.

    There are two kinds of entry:

    ``normalized``
        The extractor changed the text -- surrounding whitespace stripped, or
        spaces after a period collapsed in a descriptive field. The reviewer
        should confirm the change was safe.

    ``preserved``
        The text looks irregular but was kept exactly as published. The
        reviewer decides whether the published text is a typo worth fixing
        downstream. The extractor never makes that call on its own.
    """

    def __init__(self):
        self.entries = []

    def record(self, kind, action, table_number, location, raw, value=None,
               detail=None):
        """Record one deviation.

        Args:
            kind: What was observed, e.g. 'leading_trailing_whitespace'.
            action: Either 'normalized' or 'preserved'.
            table_number: The CH02C table the text came from, e.g. '0396'.
            location: Where in the table, e.g. 'codedContent[12].value'.
            raw: The published text, verbatim.
            value: The text as emitted, when it differs from ``raw``.
            detail: Optional extra keys to merge into the entry, e.g. how many
                separate runs of spaces a normalization collapsed.
        """
        entry = {
            'kind': kind,
            'action': action,
            'tableNumber': table_number,
            'location': location,
            'raw': raw,
        }
        if value is not None and value != raw:
            entry['normalized'] = value
        if detail:
            entry.update(detail)
        entry['section'], entry['field'] = split_location(location)
        self.entries.append(entry)

    def grouped(self):
        """Group the deviations by kind, then by which field they occurred in.

        A flat list of 1,300 whitespace notes is not reviewable. Grouping makes
        the question answerable one bucket at a time: "internal double spaces in
        code display names" is a different decision from "internal double spaces
        in prose descriptions."

        Returns:
            A list of group dicts, ordered by kind then by descending count.
        """
        buckets = {}
        for entry in self.entries:
            key = (entry['kind'], entry['action'], entry['section'], entry['field'])
            buckets.setdefault(key, []).append(entry)

        groups = []
        for (kind, action, section, field), entries in buckets.items():
            groups.append({
                'kind': kind,
                'action': action,
                'section': section,
                'field': field,
                'count': len(entries),
                'tables': sorted({e['tableNumber'] for e in entries}),
                'deviations': entries,
            })
        groups.sort(key=lambda g: (g['kind'], -g['count'], g['section'], g['field']))
        return groups


def split_location(location):
    """Split a location path into its sub-section and field name.

    ``'codedContent[12].displayName'`` becomes ``('codedContent', 'displayName')``.
    Row indices are dropped so that every row of the same column groups together.
    """
    without_index = re.sub(r'\[\d+\]', '', location)
    if '.' in without_index:
        section, field = without_index.split('.', 1)
    else:
        section, field = without_index, ''
    return section, field


def normalize_cell(raw, log, table_number, location):
    """Normalize a published cell per the fidelity policy and report anomalies.

    Two changes are made: surrounding whitespace is stripped, and in
    descriptive fields a run of two or more spaces after a period becomes one
    space. Both are recorded. Every other irregularity -- double spaces
    elsewhere in the text, non-breaking spaces, embedded newlines -- is left in
    place and reported so a reviewer can decide about it.

    The remaining-double-space check deliberately runs against the *emitted*
    value rather than the published one, so that group tracks what is still
    outstanding after the period rule rather than what was outstanding before
    it.

    Args:
        raw: The cell text exactly as python-docx read it.
        log: The DeviationLog to report into. May be None to skip reporting.
        table_number: The CH02C table number, for the report.
        location: A human-readable path to the cell, for the report.

    Returns:
        The normalized cell text.
    """
    stripped = raw.strip()
    _, field = split_location(location)
    value, collapsed_runs = normalize_descriptive_text(field, stripped)

    if log is None:
        return value

    if raw != stripped:
        log.record('leading_trailing_whitespace', 'normalized',
                   table_number, location, raw, stripped)
    if collapsed_runs:
        log.record('double_space_after_period', 'normalized',
                   table_number, location, stripped, value,
                   detail={'runsCollapsed': collapsed_runs})
    if NBSP in raw:
        log.record('non_breaking_space', 'preserved', table_number, location, raw)
    if '  ' in value:
        log.record('internal_double_space', 'preserved',
                   table_number, location, stripped, value)
    if '\n' in value:
        log.record('embedded_newline', 'preserved', table_number, location, raw)

    return value


def parse_kv_table(table, log=None, table_number=None, section=None):
    """Parse a 2-column key-value table into a dict.

    Args:
        table: The python-docx table.
        log: Optional DeviationLog for whitespace reporting.
        table_number: The CH02C table number, for the report.
        section: Name of the sub-section, e.g. 'conceptDomain', for the report.
    """
    result = {}
    for row in table.rows:
        raw_cells = [cell.text for cell in row.cells]
        if len(raw_cells) < 2:
            continue
        key = normalize_cell(raw_cells[0], log, table_number, f"{section}.<key>")
        if not key:
            continue
        value = normalize_cell(raw_cells[1], log, table_number, f"{section}.{key}")
        result[key] = value
    return result


def parse_coded_content_table(table, log=None, table_number=None):
    """Parse a coded content table (Value, Display Name, Definition, Comment, Status).

    Rows with no Value but with content in other columns are NOT discarded --
    they are returned with a ``_skipped`` marker so the caller can report them.
    Silently dropping such a row is how real published content goes missing.
    """
    if len(table.rows) < 1:
        return [], []

    # Get headers from first row
    headers = [cell.text.strip().lower() for cell in table.rows[0].cells]

    # Map header variations to canonical names
    header_map = {}
    for i, h in enumerate(headers):
        if 'value' in h and 'display' not in h:
            header_map[i] = 'value'
        elif 'display' in h:
            header_map[i] = 'displayName'
        elif 'definition' in h:
            header_map[i] = 'definition'
        elif 'comment' in h or 'usage' in h:
            header_map[i] = 'comment'
        elif 'status' in h:
            header_map[i] = 'status'
        else:
            header_map[i] = h.replace(' ', '_').replace('/', '_')

    codes = []
    skipped_rows = []
    for row_idx, row in enumerate(table.rows[1:], 1):
        raw_cells = [cell.text for cell in row.cells]

        # A row where every cell is empty is Word table padding, not content.
        if not any(c.strip() for c in raw_cells):
            continue

        entry = {}
        for i, raw in enumerate(raw_cells):
            if i not in header_map:
                continue
            field = header_map[i]
            location = f"codedContent[{row_idx}].{field}"
            entry[field] = normalize_cell(raw, log, table_number, location)

        if entry.get('value'):
            codes.append(entry)
        else:
            # Non-empty row with no code value. Never drop it silently --
            # hand it back so the caller can report it for review.
            skipped_rows.append({'rowIndex': row_idx, 'cells': entry})

    return codes, skipped_rows


def classify_table(table, log=None, table_number=None):
    """Classify a Word table by its first cell content.

    Chapter 2C is not perfectly uniform: a handful of sections start their
    Code System Version block with "Version" instead of "Effective Date", and
    one starts its Table Metadata block with "Table OID" instead of "Table".
    Those variants are recognized here so their content is not lost.

    Returns:
        A ``(kind, data)`` pair. For coded content ``data`` is itself a
        ``(codes, skipped_rows)`` pair.
    """
    if len(table.rows) < 1:
        return 'unknown', {}

    first_cell = table.rows[0].cells[0].text.strip().lower()

    def kv(section):
        return parse_kv_table(table, log, table_number, section)

    if 'concept domain name' in first_cell:
        return 'concept_domain', kv('conceptDomain')
    elif 'code system oid' in first_cell:
        return 'code_system', kv('codeSystem')
    elif first_cell in ('effective date', 'version'):
        # "Version" is the CH02C variant heading for a Code System Version block.
        return 'code_system_version', kv('codeSystemVersion')
    elif 'value set oid' in first_cell:
        return 'value_set', kv('valueSet')
    elif first_cell in ('realm',):
        return 'binding', kv('binding')
    elif first_cell in ('table', 'table oid'):
        # "Table OID" is the CH02C variant heading for a Table Metadata block.
        return 'table_metadata', kv('tableMetadata')
    elif first_cell in ('value',) or ('value' in first_cell and 'display' in table.rows[0].cells[1].text.strip().lower() if len(table.rows[0].cells) > 1 else False):
        return 'coded_content', parse_coded_content_table(table, log, table_number)
    else:
        return 'unknown', kv('unknown')


# A code table heading looks like "0685 - Item Code - External (RQD-3)".
# The four-digit form is required here because this pattern is also used to
# rescue headings that carry the wrong Word style, and a looser pattern would
# start matching ordinary body text.
CODE_TABLE_HEADING_RE = re.compile(r'^\d{4}\s*[-‐‑‒–—]\s*\S')


def is_code_table_heading(paragraph, style):
    """Decide whether a paragraph starts a code table section.

    Almost every section heading in Chapter 2C uses the "Heading 3" style.
    Two do not: 0685 (Item Code - External) and 0767 (Bolus Dose Amount Units)
    are styled "Normal" in the published document. That is a defect in the
    source -- both tables are missing from the document's own table of
    contents for the same reason -- and taking the style at face value drops
    two complete code tables and silently merges their content into the
    preceding section. So a "Normal" paragraph that otherwise looks exactly
    like a section heading is accepted too.

    Table-of-contents entries match the same text pattern, but they carry a tab
    and a page number, which is what excludes them here.
    """
    text = paragraph.text.strip()
    if not text:
        return False
    if style == 'Heading 3':
        return text[0].isdigit()
    if style == 'Normal' and '\t' not in text:
        return bool(CODE_TABLE_HEADING_RE.match(text))
    return False


def build_paragraph_index(doc):
    """Build an index mapping paragraph indices to their positions in the document XML.

    Returns a list of (paragraph_index, heading_text, table_number, table_name) for
    all Heading 3 paragraphs that represent code tables.

    Headings read "0001 - Administrative Sex", but the separator is not always a
    plain hyphen -- at least one section (0827) uses an en dash. Splitting on
    " - " alone leaves that section with the whole heading as its table number
    and no name at all, so the separator is matched as a character class.
    """
    heading_pattern = re.compile(r'^(\d[\w.]*)\s*[-‐‑‒–—]\s*(.*)$')

    table_headings = []
    for i, p in enumerate(doc.paragraphs):
        style = (p.style.name or "") if p.style else ""
        if is_code_table_heading(p, style):
            text = p.text.strip()
            match = heading_pattern.match(text)
            if match:
                table_num = match.group(1).strip()
                table_name = match.group(2).strip()
            else:
                # No separator at all -- keep the whole heading as the number
                # rather than guessing, and let the caller report it.
                table_num = text
                table_name = ''
            table_headings.append({
                'para_idx': i,
                'table_number': table_num,
                'table_name': table_name,
                'heading_text': text,
                'heading_style': style,
            })
    return table_headings


def map_tables_to_sections(doc, table_headings):
    """Map Word tables to their parent code table sections.

    Strategy: Use the XML element ordering to determine which Word tables
    fall between which heading paragraphs.
    """
    from lxml import etree

    body = doc.element.body
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Build a flat list of all body-level elements with their types
    elements = []
    para_count = 0
    table_count = 0
    for child in body:
        tag = etree.QName(child.tag).localname if '}' in child.tag else child.tag
        if tag == 'p':
            elements.append(('para', para_count))
            para_count += 1
        elif tag == 'tbl':
            elements.append(('table', table_count))
            table_count += 1
        elif tag == 'sdt':
            # Structured document tags can contain paragraphs and tables
            for sub in child.iter():
                sub_tag = etree.QName(sub.tag).localname if '}' in sub.tag else sub.tag
                if sub_tag == 'p' and sub.getparent() is not None:
                    # Only count direct paragraph children of the SDT content
                    pass  # SDT paragraphs are already counted in doc.paragraphs
                elif sub_tag == 'tbl':
                    pass  # SDT tables handled similarly

    # Build heading paragraph index set
    heading_para_indices = {h['para_idx'] for h in table_headings}

    # Walk through elements and assign tables to the most recent heading
    sections = {h['table_number']: [] for h in table_headings}
    current_section = None

    for elem_type, elem_idx in elements:
        if elem_type == 'para' and elem_idx in heading_para_indices:
            # Find which heading this is
            for h in table_headings:
                if h['para_idx'] == elem_idx:
                    current_section = h['table_number']
                    break
        elif elem_type == 'table' and current_section is not None:
            sections[current_section].append(elem_idx)

    return sections


def prune_stale_outputs(output_dir, current_table_numbers):
    """Delete output files that no longer correspond to a section in the document.

    The extractor writes one file per table, named after the table number. When
    a parsing fix changes what a table number is, the file written under the old
    name is left behind: fixing the en-dash heading of table 0827 produced a
    correct `0827.json` while a file named after its entire heading stayed in
    place from an earlier run, so the corpus held 800 files for 799 sections.

    Only files inside the generated output directory are touched, and only ones
    whose table number is not in the current document.

    Returns:
        A sorted list of the filenames that were removed.
    """
    expected = {f"{number}.json" for number in current_table_numbers}
    removed = []
    for path in sorted(Path(output_dir).glob("*.json")):
        if path.name not in expected:
            path.unlink()
            removed.append(path.name)
    return removed


def extract_all_tables(doc_path, output_dir):
    """Extract all code tables from CH02C_Tables.docx."""
    print(f"Loading {doc_path}...")
    doc = docx.Document(doc_path)
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {len(doc.tables)}")

    # Step 1: Find all table headings
    table_headings = build_paragraph_index(doc)
    print(f"  Code table sections: {len(table_headings)}")

    # Step 2: Map Word tables to their sections
    print("  Mapping tables to sections...")
    section_tables = map_tables_to_sections(doc, table_headings)

    # Step 3: Extract each section
    os.makedirs(output_dir, exist_ok=True)

    log = DeviationLog()

    index = []
    stats = {
        'total_sections': len(table_headings),
        'with_codes': 0,
        'concept_domain_only': 0,
        'metadata_only': 0,
        'total_codes': 0,
        'with_code_system': 0,
        'with_value_set': 0,
        'empty_source_tables': 0,
        'skipped_code_rows': 0,
        'errors': [],
    }

    for heading in table_headings:
        tnum = heading['table_number']
        tname = heading['table_name']
        word_table_indices = section_tables.get(tnum, [])

        record = {
            'tableNumber': tnum,
            'tableName': tname,
            'provenance': {
                'sourceFile': 'CH02C_Tables.docx',
                'headingText': heading['heading_text'],
            },
            'conceptDomain': None,
            'codeSystems': [],
            'codeSystemVersions': [],
            'valueSets': [],
            'bindings': [],
            'tableMetadata': None,
            'codedContent': [],
            'unknownTables': [],
            # Observations about the published source itself: empty tables,
            # rows without a code, duplicated blocks. Not extraction errors.
            'sourceIssues': [],
        }

        if not tname:
            record['sourceIssues'].append({
                'issue': 'heading_has_no_table_name',
                'detail': heading['heading_text'],
            })

        if heading.get('heading_style') != 'Heading 3':
            record['sourceIssues'].append({
                'issue': 'heading_has_wrong_word_style',
                'detail': f"styled {heading.get('heading_style')!r} instead of 'Heading 3'; "
                          f"this section is also missing from the document's table of contents",
            })

        for ti in word_table_indices:
            try:
                table = doc.tables[ti]

                # A table with no text at all is an empty grid in the published
                # document. Record it -- it may mean content is missing upstream.
                if not any(cell.text.strip() for row in table.rows for cell in row.cells):
                    record['sourceIssues'].append({
                        'issue': 'empty_table_in_source',
                        'wordTableIndex': ti,
                        'rows': len(table.rows),
                        'columns': len(table.columns),
                    })
                    stats['empty_source_tables'] += 1
                    continue

                ttype, data = classify_table(table, log, tnum)

                if ttype == 'concept_domain':
                    if record['conceptDomain'] is None:
                        record['conceptDomain'] = data
                    else:
                        # Duplicate concept domain (happens in ~2 tables)
                        record['conceptDomain'].update(data)
                elif ttype == 'code_system':
                    record['codeSystems'].append(data)
                elif ttype == 'code_system_version':
                    record['codeSystemVersions'].append(data)
                elif ttype == 'value_set':
                    record['valueSets'].append(data)
                elif ttype == 'binding':
                    record['bindings'].append(data)
                elif ttype == 'table_metadata':
                    if record['tableMetadata'] is None:
                        record['tableMetadata'] = data
                    else:
                        # Two metadata blocks in one section. Merge rather than
                        # overwrite so the first block's keys are not lost, and
                        # report the collision.
                        clashes = {k: [record['tableMetadata'][k], v]
                                   for k, v in data.items()
                                   if k in record['tableMetadata']
                                   and record['tableMetadata'][k] != v}
                        record['sourceIssues'].append({
                            'issue': 'duplicate_table_metadata_block',
                            'wordTableIndex': ti,
                            'conflictingKeys': clashes,
                        })
                        record['tableMetadata'].update(data)
                elif ttype == 'coded_content':
                    codes, skipped = data
                    # Append, never replace: a section split across two Word
                    # tables would otherwise lose the first half.
                    record['codedContent'].extend(codes)
                    for row in skipped:
                        record['sourceIssues'].append({
                            'issue': 'code_row_without_value',
                            'wordTableIndex': ti,
                            'rowIndex': row['rowIndex'],
                            'cells': row['cells'],
                        })
                        stats['skipped_code_rows'] += 1
                elif ttype == 'unknown':
                    record['unknownTables'].append(data)
                    record['sourceIssues'].append({
                        'issue': 'unclassified_table',
                        'wordTableIndex': ti,
                        'firstCell': table.rows[0].cells[0].text.strip()[:80],
                    })
            except Exception as e:
                stats['errors'].append({
                    'tableNumber': tnum,
                    'wordTableIndex': ti,
                    'error': str(e),
                })

        # Clean up empty lists
        if not record['unknownTables']:
            del record['unknownTables']
        if not record['sourceIssues']:
            del record['sourceIssues']

        # Update stats
        if record['codedContent']:
            stats['with_codes'] += 1
            stats['total_codes'] += len(record['codedContent'])
        elif record['tableMetadata'] and not record['codeSystems']:
            stats['metadata_only'] += 1
        else:
            stats['concept_domain_only'] += 1

        if record['codeSystems']:
            stats['with_code_system'] += 1
        if record['valueSets']:
            stats['with_value_set'] += 1

        # Write individual file
        out_path = os.path.join(output_dir, f"{tnum}.json")
        with open(out_path, 'w') as f:
            json.dump(record, f, indent=2)

        # Add to index
        index_entry = {
            'tableNumber': tnum,
            'tableName': tname,
            'hasCodeSystem': bool(record['codeSystems']),
            'hasValueSet': bool(record['valueSets']),
            'hasCodedContent': bool(record['codedContent']),
            'codeCount': len(record['codedContent']),
        }

        # Include key identifiers if available
        if record['codeSystems']:
            cs = record['codeSystems'][0]
            index_entry['codeSystemOID'] = cs.get('Code System OID', '')
            index_entry['codeSystemSymbolicName'] = cs.get('SymbolicName', '')
        if record['valueSets']:
            vs = record['valueSets'][0]
            index_entry['valueSetOID'] = vs.get('Value Set OID', '')
            index_entry['valueSetURI'] = vs.get('URI', '')
            index_entry['valueSetSymbolicName'] = vs.get('SymbolicName', '')
        if record['conceptDomain']:
            cd = record['conceptDomain']
            index_entry['conceptDomainName'] = cd.get('SymbolicName', '')
        if record['tableMetadata']:
            tm = record['tableMetadata']
            index_entry['tableType'] = tm.get('Type', '')
            index_entry['steward'] = tm.get('Steward', '')
            index_entry['tableOID'] = tm.get('Table OID', '')
        if record.get('sourceIssues'):
            index_entry['sourceIssueCount'] = len(record['sourceIssues'])

        index.append(index_entry)

    stats['pruned_files'] = prune_stale_outputs(output_dir,
                                                {h['table_number'] for h in table_headings})

    extraction_date = datetime.date.today().isoformat()

    # Write index
    index_output = {
        'extractionDate': extraction_date,
        'sourceFile': 'CH02C_Tables.docx',
        'stats': stats,
        'tables': index,
    }
    with open(INDEX_PATH, 'w') as f:
        json.dump(index_output, f, indent=2)

    # Write the deviation log -- every difference between the published text
    # and what was emitted, plus every irregularity left deliberately intact.
    # Grouped by kind and by which field it occurred in, so each bucket is one
    # reviewable decision rather than a thousand separate ones.
    groups = log.grouped()

    by_kind = {}
    by_kind_and_field = {}
    for entry in log.entries:
        by_kind[entry['kind']] = by_kind.get(entry['kind'], 0) + 1
    for group in groups:
        where = f"{group['section']}.{group['field']}" if group['field'] else group['section']
        by_kind_and_field.setdefault(group['kind'], {})[where] = group['count']

    runs_collapsed = sum(e.get('runsCollapsed', 0) for e in log.entries
                         if e['kind'] == 'double_space_after_period')

    deviations_output = {
        'extractionDate': extraction_date,
        'sourceFile': 'CH02C_Tables.docx',
        'policy': {
            'normalized': ['leading_trailing_whitespace',
                           'double_space_after_period'],
            'preserved': ['internal_double_space', 'non_breaking_space',
                          'embedded_newline'],
            'descriptiveFields': sorted(DESCRIPTIVE_FIELDS),
            'notes': {
                'double_space_after_period':
                    'ADR-0008 D3. Two or more spaces following a period are '
                    'collapsed to one, in descriptive fields only. '
                    f'{runs_collapsed} runs collapsed across '
                    f'{by_kind.get("double_space_after_period", 0)} values.',
                'internal_double_space':
                    'What REMAINS after the period rule: runs of two or more '
                    'spaces that do not follow a period, plus every run in a '
                    'field outside the descriptive set. Still outstanding.',
            },
        },
        'counts': by_kind,
        'countsByField': by_kind_and_field,
        'groups': groups,
    }
    with open(DEVIATIONS_PATH, 'w') as f:
        json.dump(deviations_output, f, indent=2)

    stats['deviations'] = by_kind
    stats['deviationsByField'] = by_kind_and_field
    return stats


def main():
    # Resolve paths relative to workspace root
    workspace = Path(__file__).resolve().parent.parent.parent
    os.chdir(workspace)

    doc_path = DOCX_PATH
    if not os.path.exists(doc_path):
        print(f"ERROR: {doc_path} not found. Expected in workspace root.")
        sys.exit(1)

    stats = extract_all_tables(doc_path, OUTPUT_DIR)

    print(f"\n=== Extraction Complete ===")
    print(f"  Total tables: {stats['total_sections']}")
    print(f"  With coded content: {stats['with_codes']}")
    print(f"  Concept domain only: {stats['concept_domain_only']}")
    print(f"  Metadata only: {stats['metadata_only']}")
    print(f"  Total codes extracted: {stats['total_codes']}")
    print(f"  With code system: {stats['with_code_system']}")
    print(f"  With value set: {stats['with_value_set']}")
    print(f"\n  Source observations (not extraction errors):")
    print(f"    Empty tables in source: {stats['empty_source_tables']}")
    print(f"    Code rows with no value: {stats['skipped_code_rows']}")
    if stats.get('pruned_files'):
        print(f"\n  Removed {len(stats['pruned_files'])} stale output file(s) "
              f"left by an earlier run:")
        for name in stats['pruned_files']:
            print(f"    {name}")
    print(f"\n  Deviations from published text (by kind, then by field):")
    by_field = stats.get('deviationsByField', {})
    for kind, count in sorted(stats.get('deviations', {}).items()):
        print(f"    {kind}: {count}")
        for where, n in sorted(by_field.get(kind, {}).items(), key=lambda kv: -kv[1]):
            print(f"        {n:5d}  {where}")
    if stats['errors']:
        print(f"  Errors: {len(stats['errors'])}")
        for err in stats['errors'][:5]:
            print(f"    Table {err['tableNumber']} (Word table {err['wordTableIndex']}): {err['error']}")
    print(f"\n  Output: {OUTPUT_DIR}/")
    print(f"  Index: {INDEX_PATH}")
    print(f"  Deviations: {DEVIATIONS_PATH}")


if __name__ == '__main__':
    main()
