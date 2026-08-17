#!/usr/bin/env python3
"""LLM-mediated V2.9.1 extraction (Phase 1 prototype, ADR-0006).

Walks an HL7 V2.9.1 source .docx chapter, identifies candidate tables (message
structures + segment field tables), and asks Claude Sonnet 4.6 to emit
JSON records matching the same schema as v291-extracted/.

The output corpus lives in v291-llm/ and is intended to be cross-validated
against the python-docx-derived v291-extracted/ corpus via
compare_python_vs_llm.py (Phase 1 stretch goal).

Cost-conscious design:
  * One LLM call per candidate table.
  * Static system prompt (schema + V2 notation rules + few-shot examples) is
    cached via prompt_caching, so the prefix amortizes across calls.
  * Per-call payload is small: the table as markdown plus a few lines of
    surrounding context.

Auth: uses Google Cloud Vertex AI via Application Default Credentials.
Reads project + region from ANTHROPIC_VERTEX_PROJECT_ID and CLOUD_ML_REGION.

Usage:
    python3 tooling/scripts/extract_v291_llm.py CH03_PatientAdmin.docx

    # Dry-run: walk the doc and report what WOULD be sent to the LLM,
    # without making API calls
    python3 tooling/scripts/extract_v291_llm.py CH03_PatientAdmin.docx --dry-run

    # Limit to first N tables for a fast sanity check
    python3 tooling/scripts/extract_v291_llm.py CH03_PatientAdmin.docx --limit 3
"""

import argparse
import itertools
import json
import os
import re
import sys
import threading
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import List, Literal

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from pydantic import BaseModel, Field


PROJECT_ROOT = Path(__file__).resolve().parents[2]
SOURCE_DIR = PROJECT_ROOT / "v2plus_docx"
OUTPUT_DIR = PROJECT_ROOT / "v291-llm"
MSG_STRUCT_OUT = OUTPUT_DIR / "message-structures"
SEGMENTS_OUT = OUTPUT_DIR / "segments"
DATA_TYPES_OUT = OUTPUT_DIR / "data-types" / "complex"
VOCABULARY_OUT = OUTPUT_DIR / "vocabulary"

MODEL = "claude-sonnet-4-6@default"
MAX_TOKENS = 8000  # plenty for one table; far less than the 64K Sonnet ceiling
# Vocabulary sections vary enormously: the median is ~1,600 characters but two
# sections exceed 65,000. A single flat max_tokens would either truncate those
# or force every small call onto the slow streaming path, so it is sized per
# section instead. See vocab_max_tokens().
VOCAB_MIN_MAX_TOKENS = 4_000
VOCAB_CEILING_MAX_TOKENS = 60_000


# ---------------------------------------------------------------------------
# Pydantic schema — mirrors v291-extracted/*.json shape
# ---------------------------------------------------------------------------


class Provenance(BaseModel):
    sourceFile: str
    chapter: str
    clause: str
    sectionHeading: str
    tableIndex: int
    captionText: str


class RawRow(BaseModel):
    segments: str
    description: str = ""
    status: str = ""
    chapter: str = ""


class ParsedElement(BaseModel):
    type: Literal["segment", "group_begin", "group_end"]
    code: str = ""
    description: str = ""
    optionality: str = ""  # "R" | "O" | "C" | "" (group markers)
    repetition: bool = False
    chapter: str = ""


class MessageStructureRecord(BaseModel):
    structureId: str
    caption: str
    provenance: Provenance
    rawRows: List[RawRow]
    parsedStructure: List[ParsedElement]


class SegmentField(BaseModel):
    sequence: str
    length: str = ""
    confLength: str = ""
    dataType: str = ""
    optionality: str = ""
    repetition: str = ""
    tableBinding: str = ""
    itemNumber: str = ""
    name: str = ""


class SegmentOccurrence(BaseModel):
    name: str
    provenance: Provenance
    fields: List[SegmentField]


class SegmentRecord(BaseModel):
    code: str
    occurrence: SegmentOccurrence


class DataTypeComponent(BaseModel):
    sequence: str
    length: str = ""
    confLength: str = ""
    dataType: str = ""
    optionality: str = ""
    tableBinding: str = ""
    name: str = ""
    comments: str = ""
    sectionRef: str = ""


class DataTypeOccurrence(BaseModel):
    name: str
    caption: str = ""
    provenance: Provenance
    components: List[DataTypeComponent]


class DataTypeRecord(BaseModel):
    code: str
    occurrence: DataTypeOccurrence


class KeyValue(BaseModel):
    """One row of a two-column metadata table in Chapter 2C.

    Modelled as an explicit key/value pair rather than a fixed set of named
    fields because Chapter 2C is not uniform: most Code System blocks use
    "URI", eight use "URL", and a few carry one-off keys like "THO URL" or
    "UTG Name". A fixed schema would silently discard those.
    """

    key: str
    value: str = ""


class KeyValueBlock(BaseModel):
    """A single two-column metadata table (one section may contain several)."""

    entries: List[KeyValue]


class CodedEntry(BaseModel):
    """One row of a Chapter 2C coded-content table."""

    value: str
    displayName: str = ""
    definition: str = ""
    comment: str = ""
    status: str = ""


class VocabularyRecord(BaseModel):
    """One Chapter 2C code table section, with all of its sub-tables."""

    tableNumber: str
    tableName: str
    conceptDomain: List[KeyValue] = []
    codeSystems: List[KeyValueBlock] = []
    codeSystemVersions: List[KeyValueBlock] = []
    valueSets: List[KeyValueBlock] = []
    bindings: List[KeyValueBlock] = []
    tableMetadata: List[KeyValue] = []
    codedContent: List[CodedEntry] = []


# Schema is split per heuristic hint rather than wrapped in a single
# discriminated union: Vertex structured_outputs rejected the union form
# as "Schema is too complex" / "Grammar compilation timed out" (2026-05-17).
# The caption-style heuristic already classifies tables reliably enough
# (CH03: 108 msg + 21 seg, 0 unknown), so the client picks the schema.
SCHEMA_FOR_HINT = {
    "message_structure": MessageStructureRecord,
    "segment": SegmentRecord,
    "data_type": DataTypeRecord,
    "vocabulary": VocabularyRecord,
}


# ---------------------------------------------------------------------------
# System prompt — static, cached
# ---------------------------------------------------------------------------


SYSTEM_PROMPT = """\
You extract structured records from HL7 V2.9.1 specification tables.

You are given:
  * The chapter heading hierarchy (clause number + section heading)
  * Two or three paragraphs of surrounding context from the source document
  * One table rendered as a Markdown table
  * A "mode" indicating which extraction shape the client wants

The client decides the mode upfront from caption styles in the source
document, so you do not need to classify the table — just extract per the
mode below.

# Mode: message_structure

Lists the segments that make up an HL7 V2 message. The leftmost column shows
segment codes wrapped in V2 cardinality notation. The other columns are
typically Description, Status, and Chapter.

V2 cardinality notation:
  * `XYZ`         — required, single (R)
  * `[ XYZ ]`     — optional, single (O)
  * `{ XYZ }`     — required, repeating (R, repetition: true)
  * `[{ XYZ }]`   — optional, repeating (O, repetition: true)
  * `[{`          — start of an optional repeating group (group_begin)
  * `}]`          — end of a group (group_end)
  * `<` `>` `|`   — choice-group markers (rare; classify as group_begin/end with code preserved)

Output shape:
{
  "structureId": "ADT_A01",   // third caret-separated token of the caption, verbatim
  "caption": "ADT^A01^ADT_A01: ADT Message",   // section heading verbatim
  "provenance": { ... fields the user provides verbatim ... },
  "rawRows": [
    // EXACT rows from the table — preserve brackets, braces, whitespace,
    // chapter numbers, status flags. This is the source-faithful view.
    { "segments": "MSH",          "description": "Message Header",   "status": "", "chapter": "2" },
    { "segments": "[{ SFT }]",    "description": "Software Segment", "status": "", "chapter": "2" }
  ],
  "parsedStructure": [
    // Interpreted view: one element per logical row.
    // Use type "segment" for actual segments, with the bare code (no brackets/braces).
    // Use type "group_begin" for "[{" rows, type "group_end" for "}]" rows.
    // For group_begin/group_end, leave `code` empty unless the source shows a name.
    { "type": "segment",     "code": "MSH", "description": "Message Header", "optionality": "R", "repetition": false, "chapter": "2" },
    { "type": "segment",     "code": "SFT", "description": "Software Segment", "optionality": "O", "repetition": true, "chapter": "2" }
  ]
}

structureId rule: take the third caret-separated token of the caption verbatim.
Do NOT synthesize per-event IDs even when the third token is not unique. For
example, the caption "ACK^R30^ACK: General Acknowledgment" yields structureId
"ACK" (not "ACK_R30"); "ADT^A01^ADT_A01: ..." yields "ADT_A01".

# Mode: segment

Lists the fields of a segment (a.k.a. "HL7 Attribute Table"). The section
heading is typically of the form "HL7 Attribute Table - PID - Patient
Identification". Columns are usually:
SEQ | LEN | C.LEN | DT | Optionality | RP/# | TBL# | ITEM# | ELEMENT NAME.

Output shape:
{
  "code": "PID",
  "occurrence": {
    "name": "Patient Identification",
    "provenance": { ... },
    "fields": [
      { "sequence": "1", "length": "1..4", "confLength": "", "dataType": "SI",
        "optionality": "O", "repetition": "", "tableBinding": "",
        "itemNumber": "00104", "name": "Set ID - PID" },
      ...
    ]
  }
}

Rules for the fields array:
  * Preserve cells verbatim. If a cell is blank, emit "".
  * `repetition` is "Y" / "" / a number, copied from the source.
  * Don't normalize, don't fix typos. Source-faithful capture is the goal.

# Mode: data_type

Lists the components of an HL7 V2 complex data type (a.k.a. "HL7 Component
Table"). The section heading is typically of the form "HL7 Component Table -
CWE – Coded with Exceptions". Columns are usually:
SEQ | LEN | C.LEN | DT | OPT | TBL# | COMPONENT NAME | COMMENTS | SEC.REF.

Note the column differences from segment mode: data types have COMMENTS and
SEC.REF. instead of segment's RP/# (repetition) and ITEM# (item number).

Output shape:
{
  "code": "CWE",
  "occurrence": {
    "name": "Coded with Exceptions",
    "caption": "HL7 Component Table - CWE – Coded with Exceptions",
    "provenance": { ... },
    "components": [
      { "sequence": "1", "length": "", "confLength": "20=", "dataType": "ST",
        "optionality": "O", "tableBinding": "", "name": "Identifier",
        "comments": "", "sectionRef": "2A.2.76" },
      ...
    ]
  }
}

Rules for the components array:
  * Preserve cells verbatim. If a cell is blank, emit "".
  * `confLength` often has trailing markers like "20=" or "199#"; copy verbatim.
  * `sectionRef` is a clause number like "2A.2.76"; copy verbatim.
  * Don't normalize, don't fix typos. Source-faithful capture is the goal.

# Mode: vocabulary

A Chapter 2C code table section. Unlike the other modes you are given SEVERAL
small tables at once — everything under one "0001 - Administrative Sex" style
heading. Each sub-table is a two-column key/value block, except the coded
content table which is a grid of codes.

Identify each sub-table by the label in its FIRST cell:

  * "Concept Domain Name"  → conceptDomain      (a key/value block)
  * "Code System OID"      → codeSystems        (may occur more than once)
  * "Effective Date" or "Version"
                           → codeSystemVersions (may occur more than once)
  * "Value Set OID"        → valueSets          (may occur more than once)
  * "Realm"                → bindings           (may occur more than once)
  * "Table" or "Table OID" → tableMetadata      (a key/value block)
  * "Value" + "Display Name" header row
                           → codedContent       (the grid of codes)

Output shape:
{
  "tableNumber": "0001",              // from the heading, digits only
  "tableName": "Administrative Sex",  // from the heading, after the dash
  "conceptDomain": [
    { "key": "Concept Domain Name", "value": "DOM: 1" },
    { "key": "SymbolicName", "value": "AdministrativeSex" },
    { "key": "Description", "value": "The domain of possible values ..." },
    { "key": "Concept Domain Only", "value": "no" }
  ],
  "codeSystems": [
    { "entries": [ { "key": "Code System OID", "value": "2.16.840.1.113883.18.458" }, ... ] }
  ],
  "codeSystemVersions": [ { "entries": [ ... ] } ],
  "valueSets": [ { "entries": [ ... ] } ],
  "bindings": [ { "entries": [ { "key": "Realm", "value": "representative" } ] } ],
  "tableMetadata": [ { "key": "Table", "value": "0001" }, ... ],
  "codedContent": [
    { "value": "F", "displayName": "Female", "definition": "", "comment": "", "status": "" },
    ...
  ]
}

Rules for vocabulary mode — fidelity matters more here than anywhere else,
because these values become published code system content:

  * Copy every key EXACTLY as printed, including capitalisation and spacing:
    "SymbolicName", "where used", "HL7 Version Introduced", "Table OID".
    Do not rename "URL" to "URI" or vice versa — some sections use each.
  * Copy every value EXACTLY as printed. Do NOT trim whitespace, collapse
    double spaces, fix spelling, expand abbreviations, or reformat dates.
    A separate pipeline compares your output against an independent extractor;
    "helpful" cleanup shows up as a false disagreement.
  * Keep the rows of each block in source order.
  * Include EVERY code row. Long code tables are common — do not summarise,
    truncate, or write "...". If a code row's Definition or Comment cell is
    long, reproduce it in full.
  * If a code row's cell is empty, emit "".
  * If a sub-table is absent from the section, emit an empty list for it.
  * If a table in the section is completely empty (no text in any cell),
    ignore it — do not invent placeholder rows.
  * status is a single letter when present ("D", "N", "B", "R") — copy it
    verbatim and never infer one that isn't printed.
  * Do NOT emit a provenance object in this mode. The client records
    provenance itself, so echoing it back would only waste output.
  * Cell text is escaped on the way in. A literal backslash-n in a cell means
    the published cell has a LINE BREAK there — emit a real newline character
    in your output at that point. A literal backslash-pipe means a "|"
    character, and a double backslash means a single backslash.

# Universal rules

  * Reproduce typographic characters EXACTLY as they appear. Curly quotes
    (“ ” ‘ ’) stay curly, en dashes (–) and em dashes (—) stay as they are,
    ellipsis characters (…) stay single characters, non-breaking spaces stay
    non-breaking. Never "helpfully" convert them to ASCII equivalents such as
    " ' - ... — the published document is the reference, and a substituted
    character reads downstream as a content change.
  * COPY the `provenance` object verbatim from the user's input — do not
    invent or modify any field.
  * Never invent rows. If a row in the source has missing cells, emit
    empty strings for those cells.
  * Never reorder rows.
  * Never collapse or expand cardinality. `[{ X }]` stays as one row in
    rawRows; in parsedStructure the X segment is one element with optionality:O
    and repetition:true.
  * If a table is clearly not what the mode says (e.g. heuristic misfired),
    extract what you can and leave fields empty where they don't apply —
    downstream tooling diffs your output against an independent extractor
    and will flag the mismatch.
"""


# ---------------------------------------------------------------------------
# Walking the docx
# ---------------------------------------------------------------------------


HEADING_STYLE_RE = re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE)
MSG_CAPTION_STYLE = "Msg Table Caption"
ATTR_CAPTION_STYLE = "Attribute Table Caption"
COMPONENT_CAPTION_STYLE = "Component Table Caption"


def iter_body(doc):
    """Yield (kind, obj) pairs in document order."""
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield "para", Paragraph(child, doc)
        elif tag == "tbl":
            yield "table", Table(child, doc)


def heading_level(p):
    style = (p.style.name or "") if p.style else ""
    m = HEADING_STYLE_RE.match(style)
    return int(m.group(1)) if m else None


def render_table_as_markdown(tbl):
    """Render a docx Table to a Markdown table. Cells with newlines are joined with ' / '."""
    rows = []
    for row in tbl.rows:
        cells = [" / ".join(p.text.strip() for p in cell.paragraphs).strip() for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]

    def esc(s):
        return s.replace("|", "\\|").replace("\n", " ")

    out = []
    out.append("| " + " | ".join(esc(c) for c in rows[0]) + " |")
    out.append("|" + "|".join(["---"] * width) + "|")
    for r in rows[1:]:
        out.append("| " + " | ".join(esc(c) for c in r) + " |")
    return "\n".join(out)


def walk_chapter(docx_path, chapter_num):
    """Walk a chapter docx and yield one tuple per table in document order.

    Tracks two parallel things alongside the body iteration:
      * Heading 1/2/3/... → numeric clause counters (mirrors the python-docx
        extractor's logic so clause numbers match across both corpuses).
      * Most recent caption-style paragraph (Msg Table Caption /
        Attribute Table Caption) → the table's logical caption.

    Yields: dict with table_idx, clause, section_heading (deepest Heading text),
    caption (caption-style text or ""), caption_style (style name or ""),
    paragraphs (last 5 body paragraphs after most recent reset), table.
    """
    doc = Document(str(docx_path))

    heading_counters = {}  # level -> counter
    deepest_heading_text = ""
    current_caption = ""
    current_caption_style = ""
    recent_paras = []
    table_index = 0

    for kind, obj in iter_body(doc):
        if kind == "para":
            style = (obj.style.name or "") if obj.style else ""
            text = obj.text.strip()
            if not text:
                continue

            level = heading_level(obj)
            if level is not None:
                if level == 1:
                    heading_counters = {}
                elif level >= 2:
                    heading_counters[level] = heading_counters.get(level, 0) + 1
                    for l in list(heading_counters):
                        if l > level:
                            del heading_counters[l]
                deepest_heading_text = text
                # Heading resets the caption + paragraph context
                current_caption = ""
                current_caption_style = ""
                recent_paras = []
            elif style in (MSG_CAPTION_STYLE, ATTR_CAPTION_STYLE, COMPONENT_CAPTION_STYLE):
                current_caption = text
                current_caption_style = style
                recent_paras = []
            else:
                recent_paras.append(text)
                if len(recent_paras) > 5:
                    recent_paras = recent_paras[-5:]
        elif kind == "table":
            # Compute clause number from active heading counters.
            parts = [chapter_num]
            if heading_counters:
                for l in range(2, max(heading_counters) + 1):
                    if l in heading_counters:
                        parts.append(str(heading_counters[l]))
            clause = ".".join(parts)
            yield {
                "table_idx": table_index,
                "clause": clause,
                "section_heading": deepest_heading_text,
                "caption": current_caption,
                "caption_style": current_caption_style,
                "paragraphs": list(recent_paras),
                "table": obj,
            }
            table_index += 1
            # Caption is consumed by the table it precedes — clear so we
            # don't re-attribute it to the next table.
            current_caption = ""
            current_caption_style = ""


VOCAB_HEADING_RE = re.compile(r'^(\d[\w.]*)\s*[-‐‑‒–—]\s*(.*)$')

# Strict four-digit form, used to rescue section headings that carry the wrong
# Word style. Kept in step with extract_v291_vocabulary.py so both pipelines
# see the same 799 sections.
VOCAB_STRICT_HEADING_RE = re.compile(r'^\d{4}\s*[-‐‑‒–—]\s*\S')


def is_vocab_section_heading(text, style):
    """Decide whether a paragraph starts a Chapter 2C code table section.

    Sections 0685 and 0767 are styled "Normal" rather than "Heading 3" in the
    published document, which also drops them from its table of contents.
    Trusting the style alone loses both tables and merges their content into
    the section above, so a "Normal" paragraph shaped exactly like a heading is
    accepted. Table-of-contents lines match the same shape but carry a tab and
    a page number, which excludes them.
    """
    if not text:
        return False
    if style == "Heading 3":
        return text[0].isdigit()
    if style == "Normal" and "\t" not in text:
        return bool(VOCAB_STRICT_HEADING_RE.match(text))
    return False


def walk_vocabulary_sections(docx_path):
    """Walk CH02C and yield one entry per code table section.

    Chapter 2C has no caption styles to key off. Instead each code table is a
    "Heading 3" paragraph like "0001 - Administrative Sex", followed by four to
    seven small tables. So the unit of extraction is the SECTION, not the
    table -- which also means ~797 LLM calls instead of ~3,791.

    Tables are attributed to the most recent code-table heading, matching what
    extract_v291_vocabulary.py does, so the two corpuses join on tableNumber.

    Yields:
        dict with table_number, table_name, heading_text, table_indices, tables.
    """
    doc = Document(str(docx_path))

    current = None
    table_index = 0

    for kind, obj in iter_body(doc):
        if kind == "para":
            style = (obj.style.name or "") if obj.style else ""
            text = obj.text.strip()
            if not is_vocab_section_heading(text, style):
                continue
            if current is not None:
                yield current
            match = VOCAB_HEADING_RE.match(text)
            current = {
                "table_number": match.group(1).strip() if match else text,
                "table_name": match.group(2).strip() if match else "",
                "heading_text": text,
                "table_indices": [],
                "tables": [],
            }
        elif kind == "table":
            if current is not None:
                current["table_indices"].append(table_index)
                current["tables"].append(obj)
            table_index += 1

    if current is not None:
        yield current


def render_vocab_table_as_markdown(tbl):
    """Render one Chapter 2C sub-table, preserving line breaks inside cells.

    The general renderer joins a cell's paragraphs with " / " and flattens
    newlines, which is fine for structural tables but destroys information
    here: 223 Chapter 2C cells contain real line breaks, and python-docx keeps
    them. If the model never sees them it cannot reproduce them, and every one
    would surface as a false disagreement between the two pipelines.

    Line breaks are therefore sent as a literal backslash-n escape, which the
    system prompt tells the model to turn back into a real newline.
    """
    rows = []
    for row in tbl.rows:
        cells = ["\n".join(p.text for p in cell.paragraphs) for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]

    def esc(s):
        return s.replace("\\", "\\\\").replace("|", "\\|").replace("\n", "\\n")

    out = ["| " + " | ".join(esc(c) for c in rows[0]) + " |",
           "|" + "|".join(["---"] * width) + "|"]
    for r in rows[1:]:
        out.append("| " + " | ".join(esc(c) for c in r) + " |")
    return "\n".join(out)


def render_section_as_markdown(entry):
    """Render every table in a code table section as one Markdown blob.

    Empty tables are skipped -- Chapter 2C contains a handful of blank grids,
    and sending them only invites the model to invent content for them.
    """
    blocks = []
    for idx, tbl in zip(entry["table_indices"], entry["tables"]):
        if not any(cell.text.strip() for row in tbl.rows for cell in row.cells):
            continue
        blocks.append(f"Sub-table {idx}:\n{render_vocab_table_as_markdown(tbl)}")
    return "\n\n".join(blocks)


def vocab_max_tokens(markdown):
    """Choose a max_tokens budget for one code table section.

    The model reproduces the section as JSON, so the output is the input text
    plus a lot of JSON punctuation -- every code row repeats five key names.
    A budget of two output tokens per four input characters proved too tight
    and truncated four code-heavy tables, so the allowance is now one token per
    input character. max_tokens is only a ceiling, not a charge, so being
    generous costs nothing except pushing the few largest sections onto the
    streaming path.
    """
    estimated_output = len(markdown)
    return max(VOCAB_MIN_MAX_TOKENS, min(estimated_output, VOCAB_CEILING_MAX_TOKENS))


def write_vocabulary_record(record, source_file, heading_text):
    """Write one vocabulary section to v291-llm/vocabulary/<tableNumber>.json.

    Provenance is added here rather than asked of the model: the client already
    knows it for certain, and not echoing it saves output tokens on every call.
    """
    VOCABULARY_OUT.mkdir(parents=True, exist_ok=True)
    payload = record.model_dump()
    payload["provenance"] = {
        "sourceFile": source_file,
        "headingText": heading_text,
    }
    path = VOCABULARY_OUT / safe_filename(record.tableNumber)
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n")
    return path


def likely_extractable(caption_style):
    """Pre-filter using the caption style. Cheap and reliable for V2.9.1 docs.

    Returns 'message_structure' | 'segment' | 'data_type' | 'unknown'.
    """
    if caption_style == MSG_CAPTION_STYLE:
        return "message_structure"
    if caption_style == ATTR_CAPTION_STYLE:
        return "segment"
    if caption_style == COMPONENT_CAPTION_STYLE:
        return "data_type"
    return "unknown"


# ---------------------------------------------------------------------------
# LLM call
# ---------------------------------------------------------------------------


# Above this many max_tokens the SDK refuses a plain request and demands
# streaming, because the call could exceed its 10-minute non-streaming budget.
# See anthropic/_base_client.py::_calculate_nonstreaming_timeout.
NONSTREAMING_MAX_TOKENS = 21_000


def call_claude(client, payload, schema, model=MODEL, max_tokens=MAX_TOKENS):
    """One LLM extraction call. `schema` is the Pydantic class for the chosen mode.

    Requests large enough to trip the SDK's non-streaming limit are sent as a
    stream instead; both paths return the same ParsedMessage shape.
    """
    user_message = (
        "Source file: {sourceFile}\n"
        "Chapter: {chapter}\n"
        "Clause: {clause}\n"
        "Section heading: {sectionHeading}\n"
        "Table index in chapter: {tableIndex}\n"
        "Caption text: {captionText}\n"
        "Mode: {hint}\n\n"
        "Surrounding paragraphs (most recent first):\n{paragraphs}\n\n"
        "Table:\n{table}\n\n"
        "Use this exact provenance object verbatim in your output:\n{provenance_json}"
    ).format(
        sourceFile=payload["provenance"]["sourceFile"],
        chapter=payload["provenance"]["chapter"],
        clause=payload["provenance"]["clause"],
        sectionHeading=payload["provenance"]["sectionHeading"],
        tableIndex=payload["provenance"]["tableIndex"],
        captionText=payload["provenance"]["captionText"],
        hint=payload["hint"],
        paragraphs="\n".join(f"  - {p}" for p in payload["paragraphs"]) or "  (none)",
        table=payload["table_md"],
        provenance_json=json.dumps(payload["provenance"], indent=2),
    )

    request = dict(
        model=model,
        max_tokens=max_tokens,
        system=[
            {
                "type": "text",
                "text": SYSTEM_PROMPT,
                "cache_control": {"type": "ephemeral"},
            }
        ],
        messages=[{"role": "user", "content": user_message}],
        output_format=schema,
    )

    if max_tokens > NONSTREAMING_MAX_TOKENS:
        with client.messages.stream(**request) as stream:
            return stream.get_final_message()
    return client.messages.parse(**request)


# ---------------------------------------------------------------------------
# Output writing
# ---------------------------------------------------------------------------


def safe_filename(*parts):
    """Build a filename from parts, replacing chars that don't play well in paths."""
    raw = "_".join(str(p) for p in parts if p)
    return re.sub(r"[^A-Za-z0-9_.-]", "_", raw) + ".json"


def write_message_structure(record):
    MSG_STRUCT_OUT.mkdir(parents=True, exist_ok=True)
    chapter = record.provenance.chapter
    fname = safe_filename(record.structureId, chapter, record.provenance.tableIndex)
    path = MSG_STRUCT_OUT / fname
    payload = record.model_dump(exclude={"kind"})
    path.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n")
    return path


def write_segment_occurrence(record, registry):
    """Append occurrence to the per-code registry; flushed at end of run."""
    registry[record.code].append(record.occurrence.model_dump())


def flush_segment_registry(registry):
    SEGMENTS_OUT.mkdir(parents=True, exist_ok=True)
    return _flush_registry(registry, SEGMENTS_OUT)


def write_data_type_occurrence(record, registry):
    registry[record.code].append(record.occurrence.model_dump())


def flush_data_type_registry(registry):
    DATA_TYPES_OUT.mkdir(parents=True, exist_ok=True)
    return _flush_registry(registry, DATA_TYPES_OUT)


def _flush_registry(registry, out_dir):
    """Merge new occurrences into per-code files, deduping by (clause, tableIndex).

    Re-running the same chapter would otherwise duplicate every occurrence. The
    dedupe key matches what the comparison script uses to join corpuses.
    """
    written = []
    for code, occurrences in registry.items():
        path = out_dir / f"{code}.json"
        if path.exists():
            existing = json.loads(path.read_text())
            existing_occurrences = existing.get("occurrences", [])
        else:
            existing_occurrences = []

        seen = set()
        merged = []
        for occ in existing_occurrences + occurrences:
            prov = occ.get("provenance", {})
            key = (prov.get("clause", ""), prov.get("tableIndex"))
            if key in seen:
                continue
            seen.add(key)
            merged.append(occ)

        path.write_text(
            json.dumps({"code": code, "occurrences": merged}, indent=2, ensure_ascii=False) + "\n"
        )
        written.append(path)
    return written


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def report_usage(total_usage):
    """Print token usage and a rough Sonnet 4.6 cost estimate."""
    print(f"\n=== Token usage ===")
    for k, v in total_usage.items():
        print(f"  {k}: {v}")
    in_cost = (total_usage["input_tokens"] * 3.0
               + total_usage["cache_creation_input_tokens"] * 3.75
               + total_usage["cache_read_input_tokens"] * 0.30) / 1_000_000
    out_cost = total_usage["output_tokens"] * 15.0 / 1_000_000
    print(f"  estimated cost: ${in_cost + out_cost:.4f} (input ${in_cost:.4f}, output ${out_cost:.4f})")


def run_vocabulary_extraction(args, docx_path):
    """Extract Chapter 2C code tables, one LLM call per code table section.

    Sections are processed concurrently. A single call takes roughly a minute
    and there are ~797 of them, so running them one at a time would take most
    of a day. The very first section is run alone to populate the prompt cache;
    fanning out before the cache exists would make every worker pay to create
    its own copy of the system prompt.
    """
    sections = list(walk_vocabulary_sections(docx_path))
    print(f"Found {len(sections)} code table sections in {args.docx}")

    if args.offset:
        sections = sections[args.offset:]
        print(f"Skipped first {args.offset} sections ({len(sections)} remaining)")
    if args.limit:
        sections = sections[: args.limit]
        print(f"Limited to first {len(sections)} sections")

    if args.resume:
        before = len(sections)
        sections = [s for s in sections
                    if not (VOCABULARY_OUT / safe_filename(s["table_number"])).exists()]
        print(f"Resume: skipped {before - len(sections)} already-extracted sections "
              f"({len(sections)} to do)")

    if args.dry_run:
        print("\n--- DRY RUN: showing first 2 sections ---")
        for s in sections[:2]:
            md = render_section_as_markdown(s)
            print(f"\nTable {s['table_number']} — {s['table_name']!r}")
            print(f"  Sub-tables: {len(s['tables'])} (word indices {s['table_indices']})")
            print("  Markdown preview:")
            print("\n".join("    " + line for line in md.splitlines()[:12]))
        total_chars = sum(len(render_section_as_markdown(s)) for s in sections)
        print(f"\n  Total markdown across {len(sections)} sections: {total_chars:,} chars"
              f" (~{total_chars // 4:,} tokens)")
        return 0

    if not sections:
        print("Nothing to do.")
        return 0

    from anthropic import AnthropicVertex
    client = AnthropicVertex()

    counts = defaultdict(int)
    total_usage = defaultdict(int)
    failures = []
    lock = threading.Lock()
    done = itertools.count(1)
    total = len(sections)

    def extract_one(section):
        """Run one section end to end. Safe to call from several threads."""
        markdown = render_section_as_markdown(section)
        payload = {
            "provenance": {
                "sourceFile": args.docx,
                "chapter": "02C",
                "clause": section["table_number"],
                "sectionHeading": section["heading_text"],
                "tableIndex": section["table_indices"][0] if section["table_indices"] else -1,
                "captionText": section["heading_text"],
            },
            "hint": "vocabulary",
            "paragraphs": [],
            "table_md": markdown,
        }
        budget = vocab_max_tokens(markdown)

        try:
            resp = call_claude(client, payload, VocabularyRecord, max_tokens=budget)
            result = resp.parsed_output
        except Exception as e:
            with lock:
                counts["error"] += 1
                failures.append({"tableNumber": section["table_number"],
                                 "error": f"{type(e).__name__}: {e}"})
                print(f"[{next(done)}/{total}] table {section['table_number']} "
                      f"FAILED: {type(e).__name__}: {e}", flush=True)
            return

        path = write_vocabulary_record(result, args.docx, section["heading_text"])

        with lock:
            u = resp.usage
            for field in ("input_tokens", "output_tokens",
                          "cache_creation_input_tokens", "cache_read_input_tokens"):
                total_usage[field] += getattr(u, field, 0) or 0
            counts["vocabulary_section"] += 1
            counts["codes"] += len(result.codedContent)
            print(f"[{next(done)}/{total}] table {section['table_number']} "
                  f"(budget {budget}) -> {len(result.codedContent)} codes, {path.name}",
                  flush=True)

    # Warm the prompt cache with one call before fanning out.
    extract_one(sections[0])
    remaining = sections[1:]

    if remaining:
        with ThreadPoolExecutor(max_workers=args.concurrency) as pool:
            list(pool.map(extract_one, remaining))

    print("\n=== Summary ===")
    for k, v in sorted(counts.items()):
        print(f"  {k}: {v}")
    if failures:
        print(f"\n  {len(failures)} sections failed -- re-run with --resume to retry just these:")
        for f in failures[:20]:
            print(f"    {f['tableNumber']}: {f['error']}")
    report_usage(total_usage)
    return 0


def main():
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("docx", help="Filename inside v2plus_docx/, e.g. CH03_PatientAdmin.docx")
    parser.add_argument("--dry-run", action="store_true", help="Walk the doc and report candidates; no LLM calls")
    parser.add_argument("--limit", type=int, default=None, help="Limit to first N candidate tables")
    parser.add_argument("--offset", type=int, default=0, help="Skip the first N candidate tables (applied before --limit)")
    parser.add_argument("--include-unknown", action="store_true",
                        help="Also send tables that the heuristic doesn't classify (more thorough, more expensive)")
    parser.add_argument("--concurrency", type=int, default=8,
                        help="Vocabulary mode: how many sections to extract at once (default 8)")
    parser.add_argument("--resume", action="store_true",
                        help="Vocabulary mode: skip sections that already have an output file")
    parser.add_argument("--mode", choices=["auto", "structural", "vocabulary"], default="auto",
                        help="Extraction shape. 'auto' picks vocabulary for CH02C_Tables.docx, "
                             "structural for every other chapter.")
    args = parser.parse_args()

    docx_path = SOURCE_DIR / args.docx
    if not docx_path.exists():
        print(f"ERROR: {docx_path} does not exist", file=sys.stderr)
        return 2

    mode = args.mode
    if mode == "auto":
        mode = "vocabulary" if args.docx.startswith("CH02C") else "structural"
    if mode == "vocabulary":
        return run_vocabulary_extraction(args, docx_path)

    # Capture an optional letter suffix (CH02A, CH04A, CH04B) — these have
    # distinct chapter codes in the python-docx corpus ("02A", clause prefix "2A").
    chapter_match = re.match(r"CH(\d+[A-Z]?)", args.docx)
    chapter_raw = chapter_match.group(1) if chapter_match else ""
    chapter = chapter_raw  # used for provenance.chapter (zero-padded form e.g. "02A")
    chapter_num = chapter_raw.lstrip("0") or "0"  # used for clause prefix e.g. "2A"

    candidates = []
    for entry in walk_chapter(docx_path, chapter_num):
        hint = likely_extractable(entry["caption_style"])
        if hint == "unknown" and not args.include_unknown:
            continue
        candidates.append({**entry, "hint": hint})

    print(f"Found {len(candidates)} candidate tables in {args.docx}")
    by_hint = defaultdict(int)
    for c in candidates:
        by_hint[c["hint"]] += 1
    for k, v in sorted(by_hint.items()):
        print(f"  {k}: {v}")

    if args.offset:
        candidates = candidates[args.offset:]
        print(f"Skipped first {args.offset} candidates ({len(candidates)} remaining)")

    if args.limit:
        candidates = candidates[: args.limit]
        print(f"Limited to first {len(candidates)} candidates")

    if args.dry_run:
        print("\n--- DRY RUN: showing first 3 candidates ---")
        for c in candidates[:3]:
            print(f"\nTable #{c['table_idx']} (clause {c['clause']}, hint={c['hint']})")
            print(f"  Section heading: {c['section_heading']}")
            print(f"  Caption: {c['caption']!r}")
            print(f"  Paragraphs preceding: {len(c['paragraphs'])}")
            md = render_table_as_markdown(c["table"])
            preview = "\n".join(md.splitlines()[:5])
            print(f"  Table preview:\n{preview}")
        return 0

    # Real extraction — Vertex AI via Application Default Credentials.
    # Picks up project_id / region from ANTHROPIC_VERTEX_PROJECT_ID / CLOUD_ML_REGION.
    from anthropic import AnthropicVertex
    client = AnthropicVertex()

    seg_registry = defaultdict(list)
    dt_registry = defaultdict(list)
    counts = defaultdict(int)
    total_usage = defaultdict(int)

    for i, c in enumerate(candidates, 1):
        # Use the caption (Msg Table Caption / Attribute Table Caption) as the
        # logical sectionHeading + captionText when available; fall back to the
        # numbered Heading text otherwise. Matches what python-docx extraction does.
        section_label = c["caption"] or c["section_heading"]
        provenance = {
            "sourceFile": args.docx,
            "chapter": chapter,
            "clause": c["clause"],
            "sectionHeading": section_label,
            "tableIndex": c["table_idx"],
            "captionText": section_label,
        }
        payload = {
            "provenance": provenance,
            "hint": c["hint"],
            "paragraphs": c["paragraphs"],
            "table_md": render_table_as_markdown(c["table"]),
        }
        print(f"[{i}/{len(candidates)}] table#{c['table_idx']} clause={c['clause']} hint={c['hint']} ... ", end="", flush=True)

        schema = SCHEMA_FOR_HINT.get(c["hint"])
        if schema is None:
            print("skip: no schema for hint (likely --include-unknown table; not supported yet)")
            counts["skipped_unknown"] += 1
            continue

        try:
            resp = call_claude(client, payload, schema)
            result = resp.parsed_output
        except Exception as e:
            print(f"FAILED: {type(e).__name__}: {e}")
            counts["error"] += 1
            continue

        u = resp.usage
        for field in ("input_tokens", "output_tokens", "cache_creation_input_tokens", "cache_read_input_tokens"):
            total_usage[field] += getattr(u, field, 0) or 0

        if c["hint"] == "message_structure":
            path = write_message_structure(result)
            print(f"msg_structure -> {path.name}")
            counts["msg_structure"] += 1
        elif c["hint"] == "segment":
            write_segment_occurrence(result, seg_registry)
            print(f"segment {result.code} (queued)")
            counts["segment_occurrence"] += 1
        elif c["hint"] == "data_type":
            write_data_type_occurrence(result, dt_registry)
            print(f"data_type {result.code} (queued)")
            counts["data_type_occurrence"] += 1

    seg_paths = flush_segment_registry(seg_registry)
    dt_paths = flush_data_type_registry(dt_registry)

    print("\n=== Summary ===")
    for k, v in sorted(counts.items()):
        print(f"  {k}: {v}")
    print(f"  segment files written: {len(seg_paths)}")
    print(f"  data type files written: {len(dt_paths)}")
    print(f"\n=== Token usage ===")
    for k, v in total_usage.items():
        print(f"  {k}: {v}")
    # rough cost estimate at Sonnet 4.6 pricing
    in_cost = (total_usage["input_tokens"] * 3.0 + total_usage["cache_creation_input_tokens"] * 3.75 + total_usage["cache_read_input_tokens"] * 0.30) / 1_000_000
    out_cost = total_usage["output_tokens"] * 15.0 / 1_000_000
    print(f"  estimated cost: ${in_cost + out_cost:.4f} (input ${in_cost:.4f}, output ${out_cost:.4f})")
    return 0


if __name__ == "__main__":
    sys.exit(main())
